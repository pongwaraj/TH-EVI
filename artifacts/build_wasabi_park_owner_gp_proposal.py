from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from th_evi.spatial import analyze_click_location


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\Wasabi_Park_Owner_GP_Proposal.docx")

SITE_NAME = "Wasabi Park"
ADDRESS = "Mahidol Rd, Tambon Nong Hoi, Mueang Chiang Mai District, Chiang Mai 50000"
LAT = 18.759897
LON = 99.017682
PROVINCE = "Chiang Mai"
START_YEAR = 2026
END_YEAR = 2035
SELL_PRICE_PER_KWH = 7.9
ELECTRICITY_COST_PER_KWH = 4.0
AVG_KWH_PER_CAR = 35.0
PROJECT_CAPEX_EX_VAT = 0
CPO_GP_RATE = 0.08
SITE_GP_SHARE_PER_KWH = 0.30
TCE_O_AND_M_PER_MONTH = 3_000

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)
GREEN = RGBColor(27, 94, 32)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def fmt_num(value: float, digits: int = 1) -> str:
    return f"{value:,.{digits}f}"


def fmt_int(value: float) -> str:
    return f"{round(value):,}"


def fmt_payback_timing(payback: tuple[int, float] | None) -> str:
    if not payback:
        return "ยังไม่คืนทุนภายในช่วงพยากรณ์"
    months_th = [
        "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
        "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม",
    ]
    year, months_into_year = payback
    month_index = min(max(int(months_into_year), 0), 11)
    return f"คาดว่าจะคืนทุนประมาณเดือน{months_th[month_index]} {year}"


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 12.5, ACCENT),
        ("Heading 3", 11.5, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | Owner GP Proposal")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Wasabi Park")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def annual_projection_rows() -> tuple[list[dict], tuple[int, float] | None]:
    rows: list[dict] = []
    cumulative = 0.0
    payback: tuple[int, float] | None = None
    for year in range(START_YEAR, END_YEAR + 1):
        result = analyze_click_location(
            lat=LAT,
            lon=LON,
            province=PROVINCE,
            year=year,
            scenario="base",
            mode="urban",
            avg_kwh_per_session=AVG_KWH_PER_CAR,
            price_per_kwh=SELL_PRICE_PER_KWH,
        )
        modeled_cars_per_day = float(result["net_sessions_per_day"])
        perception_adjusted_cars_per_day = modeled_cars_per_day
        operational_cap = 999.0
        cars_per_day = modeled_cars_per_day
        daily_kwh = cars_per_day * AVG_KWH_PER_CAR
        daily_revenue = daily_kwh * SELL_PRICE_PER_KWH
        annual_cars = cars_per_day * 365
        annual_kwh = daily_kwh * 365
        annual_revenue = daily_revenue * 365
        annual_cpo_gp = annual_revenue * CPO_GP_RATE
        annual_electricity_cost = annual_kwh * ELECTRICITY_COST_PER_KWH
        annual_site_gp_share = annual_kwh * SITE_GP_SHARE_PER_KWH
        annual_tce_o_and_m = TCE_O_AND_M_PER_MONTH * 12
        annual_contribution = annual_revenue - annual_cpo_gp - annual_electricity_cost - annual_site_gp_share - annual_tce_o_and_m
        cumulative += annual_site_gp_share
        payback_note = ""
        rows.append({
            "year": year,
            "modeled_cars_per_day": modeled_cars_per_day,
            "perception_adjusted_cars_per_day": perception_adjusted_cars_per_day,
            "operational_cap_cars_per_day": operational_cap,
            "cars_per_day": cars_per_day,
            "daily_kwh": daily_kwh,
            "daily_revenue": daily_revenue,
            "annual_cars": annual_cars,
            "annual_kwh": annual_kwh,
            "annual_revenue": annual_revenue,
            "annual_cpo_gp": annual_cpo_gp,
            "annual_electricity_cost": annual_electricity_cost,
            "annual_site_gp_share": annual_site_gp_share,
            "annual_tce_o_and_m": annual_tce_o_and_m,
            "annual_contribution": annual_contribution,
            "cumulative_contribution": cumulative,
            "utilization_pct": (daily_kwh / (180.0 * 24.0)) * 100.0,
            "note": payback_note,
        })
    return rows, payback


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis + GP Proposal")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Wasabi Park | ข้อเสนอ GP สำหรับเจ้าของโครงการ")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.6, 4.8])
    rows = [
        ("สถานที่", SITE_NAME),
        ("ที่อยู่", ADDRESS),
        ("พิกัด", f"{LAT:.6f}, {LON:.6f}"),
        ("รูปแบบความร่วมมือ", "เจ้าของโครงการไม่ลงทุน | รับ GP 0.30 บาท/kWh"),
        ("ข้อเสนอสถานี", "SINEXCEL 180 kW จำนวน 4 ช่องจอด"),
        ("ช่วงพยากรณ์", f"{START_YEAR}-{END_YEAR} รวม 10 ปี"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_executive_summary(doc: Document, current_row: dict, final_row: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    text = (
        f"Wasabi Park อยู่บนแกนถนนมหิดลฝั่งหนองหอย ซึ่งเป็นโซน gateway เชื่อมเมืองเชียงใหม่กับเส้นทางเชียงใหม่-ลำพูน และจุดนี้ผ่านเกณฑ์ความเหมาะสมของโมเดล "
        f"ในฐานะ destination site ชัดเจน สำหรับข้อเสนอฉบับนี้ เรามองจากฝั่งเจ้าของโครงการ โดยให้ผู้ลงทุนเป็นผู้ติดตั้งสถานี DC Fast Charger รุ่น SINEXCEL 180 kW จำนวน 4 ช่องจอด "
        f"และ Wasabi Park รับผลตอบแทนในรูป GP 0.30 บาทต่อ kWh โดยไม่ต้องลงทุน CAPEX เอง "
        f"ใน Base Case ปี {START_YEAR} โมเดลประเมินว่าจุดนี้มีรถ EV เข้ามาชาร์จประมาณ {fmt_num(current_row['cars_per_day'])} คัน/วัน "
        f"คิดเป็นรายได้รวมของสถานีประมาณ {fmt_int(current_row['annual_revenue'])} บาท/ปี และทำให้ Wasabi Park ได้รับ GP ประมาณ {fmt_int(current_row['annual_site_gp_share'])} บาท/ปี "
        f"จุดแข็งของไซต์คือเป็น community mall บนแนวมหิดลที่มีฐานกิจกรรมจริง และมี EVolt AC Charger 22 kW จำนวน 2 หัวชาร์จอยู่แล้ว จึงช่วยยืนยันว่ามีผู้ใช้ EV รู้จักและแวะไซต์นี้อยู่ก่อน "
        f"แม้พื้นที่จะมีคู่แข่ง EV อยู่ใกล้ แต่ด้วยข้อเสนอ 180 kW จำนวน 4 ช่องจอด ภาพลักษณ์และความมั่นใจของลูกค้าจะดีกว่ากรณีตู้เล็กอย่างมีนัยสำคัญ "
        f"ตลอดช่วงพยากรณ์ 10 ปี Wasabi Park มีโอกาสรับ GP สะสมประมาณ {fmt_int(final_row['cumulative_contribution'])} บาท หากสถานีทำยอดได้ตาม Base Case"
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def add_key_assumptions(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สมมติฐานหลักของเอกสารฉบับนี้")

    table = doc.add_table(rows=9, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [1.9, 1.3, 2.8])
    headers = ["รายการ", "ค่า", "หมายเหตุ"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("ข้อเสนอสถานี", "180 kW / 4 ช่องจอด", "SINEXCEL DC Fast Charger สำหรับยกระดับไซต์เป็น fast-charging destination"),
        ("เงินลงทุนฝั่งเจ้าของ", "0 บาท", "สมมติว่าเจ้าของโครงการไม่ลงทุน CAPEX"),
        ("กรอบเวลา", "10 ปี", "ใช้เป็นกรอบ forecast เพื่อดูรายได้ GP ที่ Wasabi Park จะได้รับ"),
        ("ดีมานด์ต่อคัน", "35 kWh/คัน", "ตีความจากโจทย์เป็นพลังงานเฉลี่ยต่อคัน"),
        ("ราคาขาย", "7.9 บาท/kWh", "สมมติฐาน Base Case เพื่อคำนวณรายได้"),
        ("ส่วนแบ่ง GP ให้ Wasabi Park", "0.30 บาท/kWh", "คำนวณจากพลังงานขายได้ของสถานี"),
        ("สถานีเดิมในไซต์", "EVolt AC 22 kW / 2 หัว", "ใช้เป็นหลักฐานว่ามีฐานผู้ใช้ EV รู้จักไซต์อยู่แล้ว"),
        ("ข้อจำกัดการวิเคราะห์", "Base Case", "ใช้ผลจาก TH-EVI โดยไม่ได้บวก upside จากแคมเปญการตลาดหรือการจับทราฟฟิกใหม่เพิ่มเติม"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_area_analysis(doc: Document, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Area Analysis")

    doc.add_paragraph(
        "จุด Wasabi Park อยู่บนแกน Mahidol-Nong Hoi ซึ่งเป็นโซนชุมชนการค้าและเส้นทางเชื่อมฝั่งเมืองกับแนวเชียงใหม่-ลำพูน "
        "โมเดลอ่านเป็น destination site และผ่านเกณฑ์ eligibility โดยมีแรงดึงจากโรงพยาบาล กิจกรรมเมืองฝั่งไนท์บาซาร์ สนามบิน และ corridor เชิงพาณิชย์ของเส้น Route 106 "
        "ในฉบับนี้เราใช้มุมมองเจ้าของโครงการ จึงโฟกัสว่าไซต์นี้สามารถสร้างทราฟฟิก EV และแปลงเป็นรายได้ GP ให้ Wasabi Park ได้มากเพียงใด หากยกระดับเป็นสถานี 180 kW จำนวน 4 ช่องจอด"
    )

    metrics = doc.add_table(rows=9, cols=3)
    metrics.style = "Table Grid"
    metrics.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(metrics, [2.0, 1.2, 2.8])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมาย"]
    for col, text in enumerate(headers):
        cell = metrics.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะพื้นที่", str(base_result["eligibility_status"]).title(), base_result["eligibility_reason"]),
        ("ความน่าสนใจเชิง GP", "น่าสนใจระดับปานกลางค่อนข้างดี", "เจ้าของไม่ต้องลงทุนเอง แต่มีโอกาสรับส่วนแบ่งตามพลังงานขายได้ของสถานี"),
        ("ประเภททำเล", str(base_result["location_type"]).title(), "โมเดลมองเป็นจุดหมายปลายทาง ไม่ใช่ highway pass-through"),
        ("ดีมานด์รวมของพื้นที่", f"{fmt_num(base_result['gross_area_demand_sessions'])} คัน/วัน", "สะท้อน demand pool ของพื้นที่รอบจุด ไม่ใช่จำนวนที่ไซต์จะรับได้ทั้งหมด"),
        ("ดีมานด์ของจุด", f"{fmt_num(base_result['net_sessions_per_day'])} คัน/วัน", "จำนวนรถ EV ที่โมเดลคาดว่าจุดนี้สามารถดึงได้ใน Base Case เมื่ออัปเกรดเป็น fast charger format ที่แข็งแรงกว่าเดิม"),
        ("พลังงานต่อวัน (Base Case)", f"{fmt_num(base_result['net_sessions_per_day'] * AVG_KWH_PER_CAR)} kWh/วัน", f"คำนวณจาก {AVG_KWH_PER_CAR:.0f} kWh/คัน"),
        ("รายได้รวมของสถานีต่อวัน", f"{fmt_int(base_result['net_sessions_per_day'] * AVG_KWH_PER_CAR * SELL_PRICE_PER_KWH)} บาท/วัน", "คำนวณที่ราคาขาย 7.9 บาท/kWh"),
        ("GP ของ Wasabi Park ต่อวัน", f"{fmt_int(base_result['net_sessions_per_day'] * AVG_KWH_PER_CAR * SITE_GP_SHARE_PER_KWH)} บาท/วัน", "คำนวณที่ GP 0.30 บาท/kWh"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = metrics.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("แรงหนุนหลักของพื้นที่")

    bullets = [
        f"POI เด่น: {', '.join(item['name'] for item in base_result['top_pois'][:5])}",
        f"Hot zone เด่น: {', '.join(item['name'] for item in base_result['top_zones'][:3])}",
        f"Business area เด่น: {', '.join(item['name'] for item in base_result['top_business_areas'][:2]) or 'ไม่มี'}",
        f"คู่แข่งใกล้เคียง: {', '.join(item['name'] for item in base_result['top_competitors'][:4])}",
        f"Nearest access anchor {fmt_num(base_result['nearest_access_anchor_km'])} กม. | nearest competitor {fmt_num(base_result['nearest_competitor_km'])} กม.",
        "ข้อสรุปเชิงพื้นที่: ทำเลนี้มี demand รองรับจริง และหากอัปเกรดเป็น 180 kW จำนวน 4 ช่องจอด จะช่วยเพิ่มความมั่นใจของลูกค้าและเหมาะกับการเก็บ GP ให้เจ้าของไซต์มากกว่ารูปแบบตู้เล็ก",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def add_financial_forecast(doc: Document, rows: list[dict], payback: tuple[int, float] | None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Base Case Forecast และรายได้ GP ของ Wasabi Park")

    note = doc.add_paragraph()
    note.add_run(
        "สูตรที่ใช้: จำนวนรถชาร์จต่อวัน = sessions/day จาก TH-EVI ใน Base Case | พลังงานต่อวัน = รถชาร์จต่อวัน x 35 kWh | "
        "รายได้รวมของสถานี = พลังงาน x 7.9 บาท/kWh | GP ของ Wasabi Park = พลังงาน x 0.30 บาท/kWh | "
        "เจ้าของไซต์ไม่ต้องลงทุน CAPEX และไม่รับภาระค่าไฟ, GP CPO หรือ O&M ในข้อเสนอฉบับนี้"
    )

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [0.55, 0.85, 0.95, 1.0, 1.0, 1.0, 1.0, 0.8])
    headers = [
        "ปี",
        "โมเดล",
        "คัน/วัน",
        "kWh/วัน",
        "รายได้/ปี",
        "GP/ปี",
        "GP สะสม",
        "Utilization",
    ]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    for row in rows:
        values = [
            str(row["year"]),
            fmt_num(row["modeled_cars_per_day"]),
            fmt_num(row["cars_per_day"]),
            fmt_num(row["daily_kwh"]),
            fmt_int(row["annual_revenue"]),
            fmt_int(row["annual_site_gp_share"]),
            fmt_int(row["cumulative_contribution"]),
            f"{fmt_num(row['utilization_pct'])}%",
        ]
        row_cells = table.add_row().cells
        for idx, value in enumerate(values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("ข้อสรุปด้านรายได้ของเจ้าของโครงการ")

    final_row = rows[-1]
    summary_rows = [
        f"ปีแรก ({START_YEAR}) คาดว่ามีรถ EV เข้ามาชาร์จประมาณ {fmt_num(rows[0]['cars_per_day'])} คัน/วัน ทำให้สถานีมีรายได้รวมประมาณ {fmt_int(rows[0]['annual_revenue'])} บาท/ปี และทำให้ Wasabi Park ได้ GP ประมาณ {fmt_int(rows[0]['annual_site_gp_share'])} บาท/ปี",
        f"ปีสุดท้ายของสัญญา ({END_YEAR}) คาดว่ามีรถเข้าชาร์จ {fmt_num(final_row['cars_per_day'])} คัน/วัน "
        f"และทำให้ Wasabi Park ได้ GP ประมาณ {fmt_int(final_row['annual_site_gp_share'])} บาท/ปี",
        f"GP สะสมตลอด 10 ปีประมาณ {fmt_int(final_row['cumulative_contribution'])} บาท",
        "สำหรับเจ้าของโครงการ จุดเด่นของข้อเสนอนี้คือมีโอกาสสร้างรายได้ประจำระยะยาวจาก GP โดยไม่ต้องเป็นผู้ลงทุนหลักของสถานี",
    ]
    for item in summary_rows:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        if "คืนทุน" in item:
            run.font.color.rgb = GREEN
            run.bold = True


def add_caveats(doc: Document, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อควรใช้ในการตัดสินใจ")

    bullets = [
        "เอกสารฉบับนี้ใช้ Base Case เพียงกรณีเดียว ตามที่ร้องขอ และใช้สมมติฐานราคาขาย 7.9 บาท/kWh ต้นทุนค่าไฟ 4.0 บาท/kWh",
        "รายได้ GP ของ Wasabi Park ในเอกสารนี้คำนวณจาก 0.30 บาทต่อ kWh ของพลังงานที่ขายได้จริง จึงเป็นรายได้ฝั่งเจ้าของไซต์โดยตรง ไม่ได้หักค่าไฟฟ้า, GP แพลตฟอร์ม CPO หรือ O&M ออกจากส่วนของเจ้าของอีกชั้น",
        "ข้อเสนอฉบับนี้สมมติว่าเจ้าของโครงการไม่ลงทุน CAPEX เอง และใช้สถานี fast charger 180 kW จำนวน 4 ช่องจอด ซึ่งให้ภาพลักษณ์และความมั่นใจดีกว่ารูปแบบตู้เล็ก",
        "ข้อมูลของโมเดลชี้ว่าพื้นที่ Wasabi Park ได้แรงหนุนจากแกน Chiang Mai-Lamphun gateway / Route 106, โรงพยาบาล และกิจกรรมเมืองฝั่ง airport-night bazaar แต่ก็มีคู่แข่งสำคัญอยู่ใกล้มาก เช่น PEA Volta Hub Chiang Mai และผู้เล่นรายใหญ่ในเมือง",
        "ข้อมูลประกอบที่มีอยู่ชี้ว่าไซต์มี EVolt AC Charger 22 kW จำนวน 2 หัวชาร์จอยู่แล้ว ซึ่งเป็นสัญญาณเชิงบวกว่ามีผู้ใช้ EV รู้จักไซต์นี้อยู่ก่อน และช่วยสนับสนุนการต่อยอดเป็นสถานี DC ที่ใหญ่ขึ้น",
        "มุมมองของเอกสารนี้จึงเน้นโอกาสของเจ้าของโครงการในการรับรายได้ประจำจาก GP โดยไม่ต้องลงทุนเอง มากกว่าการวิเคราะห์ผลตอบแทนฝั่งนักลงทุน",
        f"ระดับความเชื่อมั่นของจุดนี้ในระบบอยู่ที่ {base_result['confidence']} และคำเตือนจากระบบคือ: {', '.join(base_result['warnings']) if base_result['warnings'] else 'ไม่มี'}",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    base_result = analyze_click_location(
        lat=LAT,
        lon=LON,
        province=PROVINCE,
        year=START_YEAR,
        scenario="base",
        mode="urban",
        avg_kwh_per_session=AVG_KWH_PER_CAR,
        price_per_kwh=SELL_PRICE_PER_KWH,
    )
    rows, payback = annual_projection_rows()

    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    add_executive_summary(doc, rows[0], rows[-1])
    add_key_assumptions(doc)
    add_area_analysis(doc, base_result)
    add_financial_forecast(doc, rows, payback)
    add_caveats(doc, base_result)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
