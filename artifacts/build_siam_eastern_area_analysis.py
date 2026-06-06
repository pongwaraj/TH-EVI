from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\Siam_Eastern_Industrial_Park_Area_Analysis.docx")

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | วิเคราะห์พื้นที่สำหรับผู้บริหาร")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Siam Eastern Industrial Park")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("วิเคราะห์พื้นที่")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Siam Eastern Industrial Park")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.45, 4.85])
    rows = [
        ("สถานที่", "Siam Eastern Industrial Park"),
        ("ที่อยู่", "243V+93G, Map Yang Phon, Pluak Daeng District, Rayong 21140"),
        ("พิกัด", "13.004796, 101.144536"),
        ("ขอบเขตการวิเคราะห์", "ประเมินศักยภาพเชิงพื้นที่ของจุดในคอร์ริดอร์อุตสาหกรรมปลวกแดงสำหรับสถานีชาร์จ EV"),
        ("จัดทำเมื่อ", "6 มิถุนายน 2026"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_summary_callout(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.cell(0, 0).width = Inches(6.3)
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    text = (
        "พื้นที่ Siam Eastern Industrial Park อยู่ในคอร์ริดอร์อุตสาหกรรมปลวกแดงที่มีสัญญาณความต้องการ EV "
        "จากการเดินทางของแรงงาน โรงงาน ซัพพลายเออร์ และเส้นเชื่อม Highway 331 อยู่จริง แต่พิกัดที่ให้มายังไม่ใช่ "
        "exact site ที่โมเดลมองว่าแข็งพอสำหรับการลงทุนรอบแรก หากมองในเชิงบริหารควรตีความว่า area นี้น่าสนใจ "
        "แต่ยังควรคัดตำแหน่งย่อยที่รับ demand ได้ดีกว่าพิกัดปัจจุบัน"
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def add_key_metrics(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ตัวเลขสำคัญ")

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [2.2, 1.25, 3.05])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมายเชิงบริหาร"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะความเหมาะสม", "Low relevance", "พื้นที่มีแรงกิจกรรมจริง แต่ exact pin นี้ยังไม่ใช่จุดลงทุนลำดับแรก"),
        ("ประเภททำเลจากโมเดล", "Highway", "พื้นที่ถูกอ่านเป็นคอร์ริดอร์อุตสาหกรรม-การเดินทาง มากกว่าจะเป็น destination site"),
        ("ดีมานด์สุทธิ ณ จุด", "6.0 sessions/day", "จำนวนครั้งชาร์จที่พิกัดนี้คาดว่าจะ capture ได้จริงหลังผ่านตัวกรองการเข้าถึง"),
        ("ดีมานด์รวมของพื้นที่", "131.5 sessions/day", "สะท้อนศักยภาพของพื้นที่ปลวกแดงรอบจุดนี้ ไม่ได้หมายความว่าพิกัดนี้จะรับได้ทั้งหมด"),
        ("พลังงานต่อวัน", "168 kWh/day", "ประมาณขนาดพลังงานเชิงปฏิบัติการเบื้องต้นของ exact site นี้"),
        ("แรงจาก hot zone", "8.0 sessions/day", "แรงหนุนหลักมาจาก Highway 331 industrial connector และกลุ่มนิคมอุตสาหกรรม"),
        ("cell ใกล้เคียงที่แข็งกว่า", "Heat score 80.3", "มี cell ใกล้โซน WHA/Eastern Seaboard ที่น่าพิจารณาเป็น exact site มากกว่า"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_methodology(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("แนวคิดการใช้ข้อมูลและการอ่านผล")

    intro = (
        "ผล Area Analysis ของ TH-EVI แยกการประเมินออกเป็นสองชั้น คือ ศักยภาพของพื้นที่รอบข้าง "
        "และความเหมาะสมของพิกัด exact site เพื่อป้องกันการสรุปผิดว่า area ที่ดีต้องแปลว่าทุกแปลงใน area นั้นเหมาะลงทุนเท่ากัน"
    )
    doc.add_paragraph(intro)

    bullets = [
        ("Base demand", "ดีมานด์ตั้งต้นจากประเภททำเลและรูปแบบการเดินทางที่คาดว่าจะเกิดขึ้นในพื้นที่"),
        ("POI signal", "แรงดึงจากสถานที่สำคัญ เช่น โรงพยาบาล ศูนย์ชุมชน จุดบริการ และแหล่งงาน"),
        ("Hot zone signal", "แรงกิจกรรมเชิงพื้นที่ที่กว้างกว่า POI เดี่ยว โดยเฉพาะคอร์ริดอร์อุตสาหกรรมและเส้นทางหลัก"),
        ("Business area signal", "การตีความเชิงธุรกิจว่า area เป็น urban fringe, industrial corridor หรือ peri-urban connector"),
        ("Competitor pressure", "แรงกดจากสถานีชาร์จที่มีอยู่หรือ audit target ที่เกี่ยวข้อง"),
        ("Access and relevance filter", "ตัวกรองที่กันไม่ให้จุดที่อยู่ใน area ดี แต่เข้าถึงไม่เด่น ถูกมองว่าเป็น exact site ที่ดีเกินจริง"),
    ]
    for label, detail in bullets:
        para = doc.add_paragraph(style="List Bullet")
        lead = para.add_run(f"{label}: ")
        lead.bold = True
        lead.font.color.rgb = INK
        para.add_run(detail)


def add_site_interpretation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สิ่งที่โมเดลเห็นที่ Siam Eastern Industrial Park")

    paragraphs = [
        "โมเดลเห็นว่าพื้นที่นี้ไม่ได้อ่อน เพราะ gross area demand อยู่ระดับ 131.5 sessions/day และแรงจาก hot zone กับ business area มาจากคอร์ริดอร์อุตสาหกรรมปลวกแดงโดยตรง โดยเฉพาะ Highway 331 Pluak Daeng industrial connector และกลุ่ม WHA Eastern Seaboard",
        "อย่างไรก็ตาม exact site ที่ให้มาถูกจัดเป็น low_relevance เนื่องจาก access_ok และ urban_eligible ยังไม่ผ่าน ทำให้ demand สุทธิที่จุดนี้ถูกจำกัดไว้ที่ประมาณ 6.0 sessions/day เท่านั้น",
        "ในเชิงธุรกิจจึงควรอ่านผลนี้ว่า Siam Eastern Industrial Park อยู่ใน area ที่ควรศึกษาเชิงลึกต่อ แต่ pin ปัจจุบันยังไม่ใช่ตำแหน่งที่ดีที่สุดสำหรับการรับ demand จากโรงงาน แรงงาน และรถซัพพลายเออร์อย่างมีประสิทธิภาพ",
    ]
    for text in paragraphs:
        doc.add_paragraph(text)


def add_nearby_better_zone(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("พื้นที่ใกล้เคียงที่น่าจับตากว่า")

    doc.add_paragraph(
        "ใน heat map ของระยอง มี cell ที่แข็งกว่าใกล้เคียงอยู่แถวพิกัดประมาณ 13.077796, 101.111251 "
        "ซึ่งได้ heat score ประมาณ 80.3 และถูกหนุนโดย WHA Eastern Seaboard Industrial Estate 1, Eastern Seaboard Industrial Estate (Rayong) "
        "และ Map Yang Phon service node พร้อม business area ประเภท WHA Eastern Seaboard industrial corridor"
    )
    doc.add_paragraph(
        "สารสำคัญสำหรับผู้บริหารคือ area ปลวกแดงฝั่งนิคมอุตสาหกรรมมีศักยภาพจริง แต่ exact parcel ควรขยับเข้าใกล้ "
        "แนว frontage ที่เชื่อมโครงข่ายนิคมและ Highway 331 มากขึ้น เพื่อเพิ่มโอกาส capture demand"
    )


def add_recommendation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อสรุปเชิงบริหารและขั้นตอนถัดไป")

    recs = [
        "ให้มอง Pluak Daeng เป็น target industrial area ที่น่าศึกษาต่อสำหรับ EV charging",
        "ไม่ควรสรุปลงทุนจากพิกัด Siam Eastern Industrial Park ปัจจุบันทันที เพราะจุดนี้ยังไม่ใช่ sweet spot ของ area",
        "ใช้แนวทาง area-first, site-second คือยืนยัน area ก่อน แล้วคัด exact parcel ใหม่ในรัศมีประมาณ 3-5 กม.",
        "ให้น้ำหนักกับ frontage ที่ติดเส้นเชื่อม Highway 331, จุดรับแรงงาน, เส้นเข้าออกนิคม และจุดบริการที่มีการแวะจริง",
        "ควรทำ field audit เพิ่มเรื่องคู่แข่งจริงในแกน OR / PEA VOLTA / EleXA / MG และดูสภาพการเข้าออกของแปลงก่อนตัดสินใจเชิงพาณิชย์",
    ]
    for item in recs:
        doc.add_paragraph(item, style="List Bullet")


def add_placeholder(doc: Document, title: str, hint: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run(title)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).width = Inches(6.1)
    cell = table.cell(0, 0)
    set_cell_border(cell, color="B8C6D6", size="10")
    set_cell_shading(cell, "FAFBFC")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(title + "\n\n" + hint)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = MUTED
    cell.add_paragraph("\n\n\n")


def add_sources(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("หมายเหตุและฐานข้อมูลที่ใช้")

    items = [
        "ผลวิเคราะห์เชิงพื้นที่จาก TH-EVI local build ณ วันที่ 6 มิถุนายน 2026",
        "ใช้ข้อมูล POI, hot zones, business areas, district nodes และ competitor audit targets ของระยองชุดเริ่มต้นสำหรับโซนปลวกแดง",
        "ข้อมูลคู่แข่งของปลวกแดงในรอบนี้ยังมีส่วนที่เป็น audit target และควรยืนยันพิกัด/port mix ก่อนใช้ตัดสินใจลงทุนจริง",
        "Heat score ของ cell ใกล้เคียงใช้เพื่อชี้พื้นที่ที่ควรคัด site ต่อ ไม่ใช่ตัวเลข session ของ exact site โดยตรง",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    doc.add_paragraph()
    add_summary_callout(doc)
    doc.add_paragraph()
    add_key_metrics(doc)
    doc.add_paragraph()
    add_methodology(doc)
    doc.add_paragraph()
    add_site_interpretation(doc)
    doc.add_paragraph()
    add_nearby_better_zone(doc)
    doc.add_paragraph()
    add_recommendation(doc)

    doc.add_page_break()
    add_placeholder(doc, "ภาพที่ 1 ภาพพื้นที่ / หน้าสถานที่", "วางภาพสถานที่จริงหรือ frontage ของ Siam Eastern Industrial Park ตรงนี้")
    doc.add_paragraph()
    add_placeholder(doc, "ภาพที่ 2 แผนที่ / Heat Map", "วางภาพแผนที่ทำเลหรือ Heat Map ของพื้นที่ปลวกแดง-ระยองตรงนี้")
    doc.add_paragraph()
    add_sources(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    temp_path = OUT_PATH.with_name(OUT_PATH.stem + "_tmp.docx")
    doc.save(temp_path)
    os.replace(temp_path, OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
