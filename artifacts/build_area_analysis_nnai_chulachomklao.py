from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\Area_Analysis_CRMA_Nakhon_Nayok.docx")

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | Area Analysis สำหรับผู้บริหาร")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("โรงเรียนนายร้อย จปร. | นครนายก")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("วิเคราะห์พื้นที่")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("โรงเรียนนายร้อยพระจุลจอมเกล้า (จปร.)")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.45, 4.85])
    rows = [
        ("สถานที่", "โรงเรียนนายร้อย จปร."),
        ("ที่อยู่", "99 หมู่ที่ 1 ถนนสุวรรณศร, พรหมณี, อำเภอเมืองนครนายก, นครนายก 26000"),
        ("พิกัด", "14.281538, 101.163411"),
        ("ขอบเขตการวิเคราะห์", "ประเมินความน่าสนใจของพื้นที่สำหรับสถานีชาร์จ EV ในเชิงพื้นที่สาธารณะ"),
        ("จัดทำเมื่อ", "8 มิถุนายน 2026"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_summary_callout(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.cell(0, 0).width = Inches(6.3)
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    text = (
        "พื้นที่โรงเรียนนายร้อย จปร. ณ พิกัดที่ระบุ ยังไม่ใช่จุดเด่นสำหรับการพัฒนาสถานีชาร์จ EV สาธารณะในรอบแรก "
        "เพราะโมเดลมองว่าเป็นจุดที่มี direct access signal ค่อนข้างอ่อน และยังไม่เชื่อมกับกิจกรรมเมืองหรือ retail/service node "
        "ชัดพอ แม้จะอยู่ในจังหวัดที่มี demand จริงบางส่วน อย่างไรก็ตาม หากมองในเชิง area screening พื้นที่นครนายกโดยรวมยังมีโซนที่น่าสนใจกว่า "
        "เช่น เมืองนครนายกและบ้านนา ซึ่งเหมาะกับการคัด exact site ต่อมากกว่า"
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def add_key_metrics(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ตัวเลขสำคัญ")

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [2.25, 1.25, 3.0])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมายเชิงบริหาร"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะความเหมาะสม", "Low relevance", "พื้นที่รอบข้างมีบริบทบางส่วน แต่ exact pin นี้ยังไม่ใช่จุดที่เหมาะลงทุนก่อน"),
        ("ประเภททำเลจากโมเดล", "Highway", "ระบบมองจุดนี้เป็นแนวคอร์ริดอร์มากกว่าศูนย์กิจกรรมเมืองหรือ destination node"),
        ("ดีมานด์สุทธิ ณ จุด", "6.0 sessions/day", "จำนวนครั้งชาร์จที่พิกัดนี้คาดว่าจะ capture ได้จริงในเชิงสาธารณะยังต่ำ"),
        ("ดีมานด์รวมของพื้นที่", "28.6 sessions/day", "พื้นที่รอบข้างมี demand บางส่วน แต่ยังไม่กว้างหรือหนาพอเหมือน core zone ของจังหวัด"),
        ("พลังงานต่อวัน", "168 kWh/day", "ขนาดพลังงานเชิงปฏิบัติการโดยประมาณ หากพัฒนาที่พิกัดนี้แบบสาธารณะ"),
        ("แรงจาก hot zone", "0.0 sessions/day", "ยังไม่มีกลุ่มกิจกรรมเมืองหรือคอร์ริดอร์หลักมาหนุน exact pin นี้อย่างชัดเจน"),
        ("แรงจาก business area", "0.6 sessions/day", "ได้แรงเพียงเล็กน้อยจาก Ban Na - Pak Phli connector"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_methodology(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("แนวคิดการใช้ข้อมูลและการอ่านผล")

    intro = (
        "ผล Area Analysis ของ TH-EVI แยกการประเมินออกเป็นสองชั้น คือ ศักยภาพของพื้นที่รอบข้าง "
        "และความเหมาะสมของพิกัด exact site เพื่อไม่ให้สรุปผิดว่า area ที่พอมี demand จะเหมาะกับทุกแปลงในพื้นที่นั้นเท่ากัน"
    )
    doc.add_paragraph(intro)

    bullets = [
        ("Base demand", "ดีมานด์ตั้งต้นจากประเภททำเลและลักษณะการเดินทางที่คาดว่าจะเกิดขึ้น"),
        ("POI signal", "แรงดึงจากสถานที่สำคัญ เช่น โรงพยาบาล มหาวิทยาลัย ศูนย์ชุมชน หรือ retail"),
        ("Hot zone signal", "แรงกิจกรรมเชิงพื้นที่ เช่น city core หรือ corridor ที่มีการเดินทางหนาแน่น"),
        ("Business area signal", "มุมมองเชิงธุรกิจว่า area เป็น urban fringe, service town หรือ peri-urban connector"),
        ("Competitor pressure", "แรงกดจากสถานีชาร์จเดิมหรือ competitor audit targets"),
        ("Access and relevance filter", "ตัวกรองความเป็นไปได้ของ exact site เพื่อกันไม่ให้จุดที่อยู่ในพื้นที่ทั่วไปดูดีเกินจริง"),
    ]
    for label, detail in bullets:
        para = doc.add_paragraph(style="List Bullet")
        lead = para.add_run(f"{label}: ")
        lead.bold = True
        lead.font.color.rgb = INK
        para.add_run(detail)


def add_site_interpretation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สิ่งที่โมเดลเห็นที่โรงเรียนนายร้อย จปร.")

    paragraphs = [
        "พิกัดนี้ได้ผลลัพธ์เป็น low_relevance โดยมี net demand เพียง 6.0 sessions/day และไม่มีแรงหนุนจาก POI หรือ hot zone เด่น ๆ ณ จุดที่ระบุ แสดงว่าระบบยังไม่เห็นการ capture demand สาธารณะอย่างชัดเจนจากตำแหน่งนี้",
        "สัญญาณเดียวที่เข้ามาช่วยเล็กน้อยคือ business area ฝั่ง Ban Na - Pak Phli connector ซึ่งสะท้อนว่าพื้นที่ไม่ได้โดดเดี่ยวเต็มที่ แต่ก็ยังไม่พอให้ exact site นี้กลายเป็น public charging node ที่แข็งแรง",
        "จากชื่อสถานที่และบริบทพื้นที่ ผู้จัดทำตีความเพิ่มเติมว่าโรงเรียนนายร้อย จปร. มีลักษณะเป็นพื้นที่สถาบัน/กึ่งปิด จึงอาจไม่เหมาะกับ public charging site แบบที่ต้องพึ่งการแวะใช้งานของคนทั่วไปเป็นหลัก",
    ]
    for text in paragraphs:
        doc.add_paragraph(text)


def add_comparison(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("พื้นที่เทียบเคียงที่น่าสนใจกว่าในจังหวัด")

    doc.add_paragraph(
        "เมื่อเทียบกับพิกัดนี้ พื้นที่บ้านนากลางเมืองให้ผลดีกว่าชัดเจน โดยโมเดลอ่านว่า eligible และมี net demand "
        "ประมาณ 40.4 sessions/day จาก service-town และ gateway demand ที่ชัดกว่า"
    )
    doc.add_paragraph(
        "ขณะเดียวกัน โซนเมืองนครนายกให้ gross area demand สูงกว่ามากที่ประมาณ 398.5 sessions/day "
        "และมีแรงหนุนจาก Lotus's Nakhon Nayok, โรงพยาบาลนครนายก และ town center ซึ่งเหมาะกว่าในมุมผู้ใช้สถานีสาธารณะ"
    )
    doc.add_paragraph(
        "สำหรับ heat map ยังมี cell ใกล้เคียงที่แข็งกว่าบริเวณประมาณ 14.2190, 101.1943 ซึ่งได้ net opportunity score ราว 35.8 "
        "และนับว่าน่าสำรวจต่อมากกว่าพิกัดของโรงเรียนนายร้อย จปร."
    )


def add_recommendation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อสรุปเชิงบริหารและขั้นตอนถัดไป")

    recs = [
        "ไม่แนะนำให้ใช้พิกัดโรงเรียนนายร้อย จปร. เป็นตัวเลือกแรกสำหรับสถานีชาร์จ EV สาธารณะ",
        "หากต้องการลงทุนในนครนายก ควรให้น้ำหนักกับโซนเมืองนครนายกหรือบ้านนา ซึ่งมี city/service demand ชัดกว่า",
        "หากโจทย์จริงของพื้นที่นี้เป็นการใช้งานภายในองค์กร เช่น รถราชการ รถบริการ หรือ fleet เฉพาะกิจ ควรแยกเป็น captive charging case ไม่ใช่ public charging case",
        "ควรทำ field audit ต่อในแนวบ้านนา เมืองนครนายก และ frontage ที่รับรถผ่านจริง เพื่อคัด exact parcel ที่เข้าถึงง่ายกว่า",
        "ควรเก็บ competitor จริงในจังหวัดเพิ่ม โดยเฉพาะรอบเมืองนครนายกและแนว Route 305/Route 33 เพื่อทำ opportunity screening ให้คมขึ้น",
    ]
    for item in recs:
        doc.add_paragraph(item, style="List Bullet")


def add_placeholder(doc: Document, title: str, hint: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run(title)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).width = Inches(6.1)
    cell = table.cell(0, 0)
    set_cell_border(cell, color="B8C6D6", size="10")
    set_cell_shading(cell, "FAFBFC")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(title + "\n\n" + hint)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = MUTED
    cell.add_paragraph("\n\n\n")


def add_sources(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("หมายเหตุและฐานข้อมูลที่ใช้")

    items = [
        "ผลวิเคราะห์เชิงพื้นที่จาก TH-EVI local build ณ วันที่ 8 มิถุนายน 2026",
        "ใช้ข้อมูล POI, hot zones, business areas, district nodes และ competitor audit targets ของนครนายกชุดเริ่มต้น",
        "ตัวเลขนี้ใช้เพื่อคัดกรองเชิงพื้นที่และเปรียบเทียบ area ไม่ควรตีความเป็น forecast เชิงพาณิชย์สุดท้ายของพิกัดโดยตรง",
        "ข้อความที่ระบุว่าพื้นที่มีลักษณะกึ่งปิด/สถาบัน เป็นข้อสังเกตเชิงบริบทจากประเภทสถานที่และควรยืนยันกับสภาพการเข้าถึงจริงอีกครั้ง",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    doc.add_paragraph()
    add_summary_callout(doc)
    doc.add_paragraph()
    add_key_metrics(doc)
    doc.add_paragraph()
    add_methodology(doc)
    doc.add_paragraph()
    add_site_interpretation(doc)
    doc.add_paragraph()
    add_comparison(doc)
    doc.add_paragraph()
    add_recommendation(doc)

    doc.add_page_break()
    add_placeholder(doc, "ภาพที่ 1 ภาพพื้นที่ / ทางเข้าโครงการ", "วางภาพพื้นที่จริงหรือ frontage การเข้าออกของโรงเรียนนายร้อย จปร. ตรงนี้")
    doc.add_paragraph()
    add_placeholder(doc, "ภาพที่ 2 แผนที่ / Heat Map", "วางภาพแผนที่ทำเลหรือ Heat Map ของนครนายกเพื่อเปรียบเทียบกับเมืองนครนายกและบ้านนา")
    doc.add_paragraph()
    add_sources(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    temp_path = OUT_PATH.with_name(OUT_PATH.stem + "_tmp.docx")
    doc.save(temp_path)
    os.replace(temp_path, OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
