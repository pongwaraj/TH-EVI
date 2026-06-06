from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\K-CITY_BANGPU_Area_Analysis.docx")

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | วิเคราะห์พื้นที่สำหรับผู้บริหาร")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("K-CITY BANGPU")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("วิเคราะห์พื้นที่")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("K-CITY BANGPU")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.45, 4.85])
    rows = [
        ("สถานที่", "K-CITY BANGPU"),
        ("ที่อยู่", "309 Sukhumvit Rd, Bang Pu, Mueang Samut Prakan District, Samut Prakan 10280"),
        ("พิกัด", "13.494853, 100.775520"),
        ("ขอบเขตการวิเคราะห์", "ประเมินศักยภาพเชิงพื้นที่สำหรับสถานีชาร์จ EV และความเหมาะสมของทำเล"),
        ("จัดทำเมื่อ", "6 มิถุนายน 2026"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_summary_callout(doc: Document) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.cell(0, 0).width = Inches(6.3)
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    text = (
        "พื้นที่บางปูมีสัญญาณความต้องการใช้สถานีชาร์จ EV อยู่จริง แต่พิกัด K-CITY BANGPU ที่เสนอมา "
        "ยังไม่ใช่จุดที่โมเดลปัจจุบันมองว่าเหมาะจะเป็นตัวเลือกแรกสำหรับการลงทุนตั้งสถานีชาร์จ "
        "กล่าวคือ พื้นที่รอบข้างมี activity และมี demand pool ที่วัดได้ แต่ตัวพิกัดนี้เองยังอ่อนในมุมการเข้าถึงโดยตรง "
        "และแรงรองรับจากบริบทเมืองโดยรอบ"
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def add_key_metrics(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ตัวเลขสำคัญ")

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [2.2, 1.2, 3.1])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมายเชิงบริหาร"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะความเหมาะสม", "Low relevance", "พื้นที่มี activity แต่พิกัดนี้ยังไม่แข็งพอสำหรับเป็นจุดลงทุนลำดับแรก"),
        ("ดีมานด์สุทธิ ณ จุด", "6.0 sessions/day", "จำนวนครั้งชาร์จที่คาดว่าจุดนี้จะ capture ได้จริงหลังผ่านตัวกรองด้านการเข้าถึงและความเกี่ยวข้อง"),
        ("ดีมานด์รวมของพื้นที่", "37.9 sessions/day", "สะท้อนว่าพื้นที่บางปูโดยรอบมี demand pool อยู่จริง"),
        ("พลังงานต่อวัน", "192 kWh/day", "ใช้เพื่อดูขนาดเชิงคร่าว ๆ ของพลังงาน ไม่ใช่ข้อสรุปสุดท้ายของขนาดสถานี"),
        ("Heat cell ที่แข็งกว่าใกล้เคียง", "38.4 sessions/day", "มี cell ใกล้เคียงที่ดูแข็งกว่าพิกัด K-CITY สำหรับการคัด exact site"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_methodology(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("แนวคิดการอ่านผลวิเคราะห์")

    intro = (
        "โมเดลไม่ได้มองพิกัดเพียงจุดเดียวแบบโดด ๆ แต่จะรวมความต้องการเชิงจุดเข้ากับ activity field ของพื้นที่รอบข้าง "
        "เพื่อแยกให้ออกระหว่าง พื้นที่ที่น่าสนใจ กับพิกัดที่ยังไม่เหมาะเป็น exact site"
    )
    doc.add_paragraph(intro)

    bullets = [
        ("Base demand", "ดีมานด์ตั้งต้นจากการเดินทางและบริบทพื้นฐานของทำเล"),
        ("POI signal", "แรงดึงจากจุดสำคัญใกล้เคียง เช่น ขนส่ง โรงพยาบาล ค้าปลีก ศูนย์ชุมชน ท่องเที่ยว หรือแหล่งงาน"),
        ("Hot zone signal", "กิจกรรมรวมระดับพื้นที่ที่สะท้อนแรงธุรกิจเกินกว่าจะมองจาก POI เดียว"),
        ("Business area signal", "การตีความพื้นที่เชิงธุรกิจ เช่น industrial connector, commuter band หรือ destination fringe"),
        ("Competitor pressure", "แรงกดจากสถานีชาร์จเดิม เมื่อมีพิกัดยืนยันที่เชื่อถือได้"),
        ("Access and relevance filter", "ตัวกรองขั้นสุดท้ายที่กันไม่ให้พื้นที่ที่ดู active ตามทฤษฎี ถูกมองว่าเป็น site ที่ดี ทั้งที่การเข้าถึงจริงยังอ่อน"),
    ]
    for label, detail in bullets:
        para = doc.add_paragraph(style="List Bullet")
        lead = para.add_run(f"{label}: ")
        lead.bold = True
        lead.font.color.rgb = INK
        para.add_run(detail)


def add_site_interpretation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สิ่งที่โมเดลเห็นที่ K-CITY BANGPU")

    paragraphs = [
        "ที่พิกัดที่ส่งมา โมเดลจัดจุดนี้เป็น road-access candidate มากกว่าจะเป็น urban charging location ที่มีแรงรองรับชัดเจน ในเชิงธุรกิจแปลว่า พื้นที่ไม่ได้ว่างเปล่า แต่ frontage ของจุดนี้เองยังไม่ใช่ตำแหน่งที่ดีที่สุดสำหรับการ capture demand ได้สม่ำเสมอทุกวัน",
        "แรงดีมานด์รอบจุดนี้ถูกดึงหลัก ๆ จาก Bang Pu Recreation Center, โซน Bang Pu industrial and coastal corridor และ business area ประเภท Bang Pu coastal destination and civic fringe ซึ่งเป็น positive signals จริง แต่ยังไม่กระจุกตัวพอที่ exact pin นี้จนจะสรุปว่าพร้อมลงทุนได้ทันที",
        "ภาพของคู่แข่งฝั่งสมุทรปราการตะวันออกยังไม่สมบูรณ์ เพราะหลายจุดยังเป็น audit target ที่ไม่มีพิกัดยืนยัน ดังนั้นความเสี่ยงด้านการแข่งขันในผลปัจจุบันมีโอกาสถูกประเมินต่ำกว่าความเป็นจริง",
    ]
    for text in paragraphs:
        doc.add_paragraph(text)


def add_recommendation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อสรุปเชิงบริหารและขั้นตอนถัดไป")

    recs = [
        "ให้มองบางปูเป็น target area ที่ควรศึกษาต่อ ไม่ใช่พื้นที่ที่ถูกตัดทิ้ง",
        "ไม่ควรล็อกการตัดสินใจจากพิกัด K-CITY ปัจจุบันเพียงจุดเดียว",
        "ใช้แนวทาง area-first, site-second คือยืนยันโซนก่อน แล้วคัด exact parcel ในรัศมีประมาณ 2-5 กม.",
        "ให้น้ำหนักกับแปลงที่ติดสุขุมวิทชัดขึ้น เข้าถึงง่ายขึ้น หรือเชื่อมกับ industrial และ municipal service core ได้ดีกว่า",
        "ควรทำ field audit เพิ่มทั้งเรื่องพิกัดคู่แข่งและสภาพการเข้าถึงของไซต์ ก่อนตัดสินใจเชิงพาณิชย์",
    ]
    for item in recs:
        doc.add_paragraph(item, style="List Bullet")


def add_placeholder(doc: Document, title: str, hint: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run(title)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).width = Inches(6.1)
    cell = table.cell(0, 0)
    set_cell_border(cell, color="B8C6D6", size="10")
    set_cell_shading(cell, "FAFBFC")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(title + "\n\n" + hint)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = MUTED
    cell.add_paragraph("\n\n\n")


def add_sources(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("หมายเหตุและฐานข้อมูลที่ใช้")

    items = [
        "ผลวิเคราะห์เชิงพื้นที่จาก TH-EVI local build ณ วันที่ 6 มิถุนายน 2026",
        "ข้อมูลสนับสนุนของสมุทรปราการครอบคลุม demand ระดับจังหวัด, บริบทประชากรรายอำเภอ, POI เฉพาะบางปู, hot zones, business areas และ competitor audit targets เบื้องต้น",
        "ข้อมูลคู่แข่งในฝั่งสมุทรปราการตะวันออกยังเป็น seed บางส่วน และควร field-verify ก่อนนำไปใช้ตัดสินใจลงทุน",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    doc.add_paragraph()
    add_summary_callout(doc)
    doc.add_paragraph()
    add_key_metrics(doc)
    doc.add_paragraph()
    add_methodology(doc)
    doc.add_paragraph()
    add_site_interpretation(doc)
    doc.add_paragraph()
    add_recommendation(doc)

    doc.add_page_break()
    add_placeholder(doc, "ภาพที่ 1 หน้าร้าน / ภาพพื้นที่จริง", "วางภาพหน้าร้านหรือภาพ frontage ของพื้นที่ตรงนี้")
    doc.add_paragraph()
    add_placeholder(doc, "ภาพที่ 2 แผนที่ / ภาพ Area Demand", "วางภาพ Heat Map, แผนที่ corridor หรือแผนที่ annotated ตรงนี้")
    doc.add_paragraph()
    add_sources(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    temp_path = OUT_PATH.with_name(OUT_PATH.stem + "_tmp.docx")
    doc.save(temp_path)
    os.replace(temp_path, OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
