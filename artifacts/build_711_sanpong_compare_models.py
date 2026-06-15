from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from th_evi.spatial import analyze_click_location


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\7-11_Sanpong_Compare_Models.docx")

SITE_NAME = "7-11 สันโป่ง"
ADDRESS = "San Pong, Mae Rim District, Chiang Mai 50180"
LAT = 18.942782
LON = 98.942951
PROVINCE = "Chiang Mai"
START_YEAR = 2026
END_YEAR = 2035
SELL_PRICE_PER_KWH = 7.9
ELECTRICITY_COST_PER_KWH = 4.0
CPO_GP_RATE = 0.08
OWNER_GP_PER_KWH = 0.25
AVG_KWH_PER_CAR = 35.0
OWNER_CAPEX_EX_VAT = 1_500_000
OWNER_O_AND_M_PER_YEAR = 36_000
RECOMMENDED_POWER_KW = 120
RECOMMENDED_CONNECTORS = 2
SESSION_CAP_PER_DAY = 18.0

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)
GREEN = RGBColor(27, 94, 32)
AMBER = RGBColor(180, 83, 9)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def fmt_num(value: float, digits: int = 1) -> str:
    return f"{value:,.{digits}f}"


def fmt_int(value: float) -> str:
    return f"{round(value):,}"


def fmt_payback_years(payback: tuple[int, float] | None) -> str:
    if not payback:
        return "ยังไม่คืนทุนภายในช่วงพยากรณ์"
    years = (payback[0] - START_YEAR) + (payback[1] / 12.0)
    return f"ประมาณ {years:.1f} ปี"


def fmt_payback_month(payback: tuple[int, float] | None) -> str:
    if not payback:
        return "ยังไม่คืนทุนภายในช่วงพยากรณ์"
    months_th = [
        "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
        "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม",
    ]
    year, month_in_year = payback
    month_index = min(max(int(month_in_year), 0), 11)
    return f"{months_th[month_index]} {year}"


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 12.5, ACCENT),
        ("Heading 3", 11.5, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | Compare Models")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(SITE_NAME)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def projection_rows() -> tuple[list[dict], dict, tuple[int, float] | None]:
    rows: list[dict] = []
    base_result = analyze_click_location(
        lat=LAT,
        lon=LON,
        province=PROVINCE,
        year=START_YEAR,
        scenario="base",
        mode="urban",
        avg_kwh_per_session=AVG_KWH_PER_CAR,
        price_per_kwh=SELL_PRICE_PER_KWH,
    )
    base_gross = float(base_result["gross_area_demand_sessions"])
    base_sessions = float(base_result["net_sessions_per_day"])
    owner_cumulative = 0.0
    gp_cumulative = 0.0
    payback: tuple[int, float] | None = None

    for year in range(START_YEAR, END_YEAR + 1):
        result = analyze_click_location(
            lat=LAT,
            lon=LON,
            province=PROVINCE,
            year=year,
            scenario="base",
            mode="urban",
            avg_kwh_per_session=AVG_KWH_PER_CAR,
            price_per_kwh=SELL_PRICE_PER_KWH,
        )
        gross_area = float(result["gross_area_demand_sessions"])
        grown_sessions = base_sessions * (gross_area / max(base_gross, 0.1))
        sessions = min(grown_sessions, SESSION_CAP_PER_DAY)
        daily_kwh = sessions * AVG_KWH_PER_CAR
        annual_kwh = daily_kwh * 365
        annual_revenue = annual_kwh * SELL_PRICE_PER_KWH
        annual_cpo_gp = annual_revenue * CPO_GP_RATE
        annual_electricity = annual_kwh * ELECTRICITY_COST_PER_KWH
        annual_owner_cf = annual_revenue - annual_cpo_gp - annual_electricity - OWNER_O_AND_M_PER_YEAR
        annual_owner_gp = annual_kwh * OWNER_GP_PER_KWH

        prev_owner_cumulative = owner_cumulative
        owner_cumulative += annual_owner_cf
        gp_cumulative += annual_owner_gp
        if payback is None and prev_owner_cumulative < OWNER_CAPEX_EX_VAT <= owner_cumulative:
            months = ((OWNER_CAPEX_EX_VAT - prev_owner_cumulative) / max(annual_owner_cf, 1.0)) * 12.0
            payback = (year, months)

        rows.append({
            "year": year,
            "gross_area": gross_area,
            "sessions_per_day": sessions,
            "daily_kwh": daily_kwh,
            "annual_revenue": annual_revenue,
            "annual_cpo_gp": annual_cpo_gp,
            "annual_electricity": annual_electricity,
            "annual_owner_cf": annual_owner_cf,
            "owner_cumulative": owner_cumulative,
            "annual_owner_gp": annual_owner_gp,
            "gp_cumulative": gp_cumulative,
            "utilization_pct": (daily_kwh / (RECOMMENDED_POWER_KW * 24.0)) * 100.0,
        })

    return rows, base_result, payback


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis + Compare Models")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("7-11 สันโป่ง | เปรียบเทียบเจ้าของลงทุนเอง vs รับ GP")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.8, 4.6])
    rows = [
        ("สถานที่", SITE_NAME),
        ("ที่อยู่", ADDRESS),
        ("พิกัด", f"{LAT:.6f}, {LON:.6f}"),
        ("รูปแบบเปรียบเทียบ", "1) เจ้าของลงทุนเอง CAPEX 1.5 ล้านบาท  2) เจ้าของรับ GP 0.25 บาท/kWh"),
        ("ขนาดสถานีอ้างอิง", f"SINEXCEL {RECOMMENDED_POWER_KW} kW | {RECOMMENDED_CONNECTORS} หัวชาร์จ"),
        ("ช่วงพยากรณ์", f"{START_YEAR}-{END_YEAR} รวม 10 ปี"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_executive_summary(doc: Document, first_row: dict, final_row: dict, base_result: dict, payback: tuple[int, float] | None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    text = (
        f"จุด 7-11 สันโป่งอยู่ในแนวแม่ริม-โชตนา-Route 107 ซึ่งมีดีมานด์เชิงคอร์ริดอร์อยู่จริง แต่ direct access signal ของจุดยังไม่แข็ง "
        f"โมเดลจึงจัดพื้นที่นี้เป็น low relevance และให้ดีมานด์ตั้งต้นเพียงประมาณ {fmt_num(base_result.get('net_sessions_per_day', 0))} คัน/วัน "
        f"เพื่อให้เห็นภาพธุรกิจในระยะ 10 ปีอย่างเป็นรูปธรรม เอกสารฉบับนี้ใช้ Base Case แบบเติบโตตาม area-demand growth ของพื้นที่ และยังคุมไม่ให้เกินศักยภาพของสถานี {RECOMMENDED_POWER_KW} kW {RECOMMENDED_CONNECTORS} หัวชาร์จ "
        f"ภายใต้กรอบนี้ ปีแรกคาดว่ามีรถ EV เข้ามาชาร์จประมาณ {fmt_num(first_row['sessions_per_day'])} คัน/วัน "
        f"ถ้าเจ้าของลงทุนเอง CAPEX {fmt_int(OWNER_CAPEX_EX_VAT)} บาท จะมีกระแสเงินสดสุทธิปีแรกประมาณ {fmt_int(first_row['annual_owner_cf'])} บาท/ปี "
        f"และคาดว่าจะคืนทุนประมาณ {fmt_payback_years(payback)} หรือราว {fmt_payback_month(payback)} "
        f"ในอีกทางหนึ่ง หากเจ้าของไม่ลงทุนเองและรับ GP {OWNER_GP_PER_KWH:.2f} บาท/kWh ปีแรกจะได้ GP ประมาณ {fmt_int(first_row['annual_owner_gp'])} บาท/ปี และ GP สะสม 10 ปีประมาณ {fmt_int(final_row['gp_cumulative'])} บาท "
        f"ดังนั้นสำหรับจุดนี้ โมเดลมองว่าแบบ GP มีความเสี่ยงต่ำกว่าอย่างชัดเจน ขณะที่แบบลงทุนเองยังคุ้มได้แต่ต้องยอมรับ horizon ที่ยาวกว่า"
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def add_key_assumptions(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สมมติฐานหลัก")

    table = doc.add_table(rows=10, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [1.9, 1.4, 2.7])
    headers = ["รายการ", "ค่า", "หมายเหตุ"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("กรอบเวลา", "10 ปี", "ใช้ดูภาพรายได้ระยะกลางของ 2 โมเดล"),
        ("สถานีอ้างอิง", f"{RECOMMENDED_POWER_KW} kW / {RECOMMENDED_CONNECTORS} หัว", "ขนาดที่เหมาะสมกับจุดนี้มากกว่าเริ่มด้วยสถานีใหญ่เกินไป"),
        ("Capex เจ้าของลงทุนเอง", "1.50 ล้านบาท", "ไม่รวม VAT"),
        ("ราคาขาย", "7.9 บาท/kWh", "ใช้คำนวณรายได้รวมของสถานี"),
        ("ต้นทุนค่าไฟ", "4.0 บาท/kWh", "ใช้ในโมเดลเจ้าของลงทุนเอง"),
        ("CPO GP", "8% ของรายได้", "ใช้ในโมเดลเจ้าของลงทุนเอง"),
        ("O&M", "36,000 บาท/ปี", "ใช้ในโมเดลเจ้าของลงทุนเอง"),
        ("GP เจ้าของพื้นที่", "0.25 บาท/kWh", "ใช้ในโมเดลรับ GP โดยไม่ลงทุน"),
        ("การเติบโตตามปี", "อิง area-demand growth", "ใช้ดีมานด์พื้นที่ของ TH-EVI เป็นตัวพา sessions/day ให้โตตามปี"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_area_analysis(doc: Document, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Area Analysis")

    doc.add_paragraph(
        "จุด 7-11 สันโป่งอยู่ในแนวแม่ริม ซึ่งมีแรงจาก hot zone Mae Rim / Chotana / Route 107 และ node ของอำเภอแม่ริม "
        "อย่างไรก็ดี จุดนี้ยังไม่ติด anchor เมืองหรือ commercial node ที่หนักพอ โมเดลจึงให้ผลเป็น low relevance และควรเริ่มจากข้อเสนอที่ประหยัดการลงทุนก่อน"
    )

    metrics = doc.add_table(rows=8, cols=3)
    metrics.style = "Table Grid"
    metrics.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(metrics, [2.0, 1.2, 2.8])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมาย"]
    for col, text in enumerate(headers):
        cell = metrics.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะพื้นที่", str(base_result.get("eligibility_status", "-")).title(), str(base_result.get("eligibility_reason", "-"))),
        ("ประเภททำเล", str(base_result.get("location_type", "-")).title(), "โมเดลอ่านเป็น highway-side opportunity มากกว่า urban node"),
        ("ดีมานด์รวมของพื้นที่", f"{fmt_num(base_result.get('gross_area_demand_sessions', 0))} คัน/วัน", "สะท้อน demand pool รอบกว้างของแนวแม่ริม"),
        ("ดีมานด์ของจุด", f"{fmt_num(base_result.get('net_sessions_per_day', 0))} คัน/วัน", "ค่า point estimate โดยตรงของโมเดลในปีแรก"),
        ("Hot zone หลัก", ", ".join(item['name'] for item in base_result.get('top_zones', [])[:2]) or "-", "ตัวผลักหลักของจุดนี้มาจาก corridor ไม่ใช่ POI เมือง"),
        ("คู่แข่งใกล้สุด", f"{fmt_num(base_result.get('nearest_competitor_km', 0))} กม.", "รอบจุดไม่ได้ถูกกดด้วย competitor ใกล้มาก แต่ฝั่ง access ยังอ่อน"),
        ("คำแนะนำขนาดสถานี", f"{RECOMMENDED_POWER_KW} kW / {RECOMMENDED_CONNECTORS} หัว", "เหมาะกว่าเริ่มด้วย 180 kW 4 ช่องจอดสำหรับไซต์นี้"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = metrics.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_comparison_table(doc: Document, rows: list[dict]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("เปรียบเทียบ 2 โมเดล")

    note = doc.add_paragraph()
    note.add_run(
        "ตารางนี้ใช้จำนวนรถ EV ต่อวันชุดเดียวกันทั้ง 2 โมเดล เพื่อให้เปรียบเทียบได้ตรง: "
        "1) เจ้าของลงทุนเองและรับรายได้สถานีหลังหักค่าไฟ, CPO GP และ O&M "
        "2) เจ้าของไม่ลงทุน แต่รับ GP 0.25 บาท/kWh"
    )

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [0.55, 0.80, 0.90, 1.05, 1.05, 1.05, 1.00, 1.00])
    headers = ["ปี", "คัน/วัน", "kWh/วัน", "รายได้/ปี", "ลงทุนเอง CF/ปี", "ลงทุนเองสะสม", "GP/ปี", "GP สะสม"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    for row in rows:
        values = [
            str(row["year"]),
            fmt_num(row["sessions_per_day"]),
            fmt_num(row["daily_kwh"]),
            fmt_int(row["annual_revenue"]),
            fmt_int(row["annual_owner_cf"]),
            fmt_int(row["owner_cumulative"]),
            fmt_int(row["annual_owner_gp"]),
            fmt_int(row["gp_cumulative"]),
        ]
        row_cells = table.add_row().cells
        for idx, value in enumerate(values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK


def add_recommendation(doc: Document, rows: list[dict], payback: tuple[int, float] | None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อเสนอแนะ")

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = callout.cell(0, 0)
    set_cell_border(cell, color="E8B26A", size="10")
    set_cell_shading(cell, "FFF7ED")
    p = cell.paragraphs[0]
    run = p.add_run(
        f"ข้อเสนอที่เหมาะสมกับ 7-11 สันโป่ง: ถ้าเจ้าของพื้นที่ต้องการความเสี่ยงต่ำ ควรเลือกโมเดลรับ GP 0.25 บาท/kWh เพราะไม่ต้องใช้เงินลงทุนเอง "
        f"ส่วนโมเดลลงทุนเอง CAPEX {fmt_int(OWNER_CAPEX_EX_VAT)} บาท แม้คาดว่าคืนทุนได้ราว {fmt_payback_years(payback)} แต่เป็น horizon ที่ค่อนข้างยาวสำหรับจุดที่ยังถูกจัดเป็น low relevance"
    )
    run.font.color.rgb = AMBER
    run.bold = True

    bullets = [
        f"ขนาดสถานีที่แนะนำยังเป็น {RECOMMENDED_POWER_KW} kW {RECOMMENDED_CONNECTORS} หัวชาร์จ ไม่ควรเริ่มด้วยสถานีใหญ่กว่านี้",
        f"ถ้าในอนาคตมีหลักฐานทราฟฟิก EV หน้างานจริงสูงเกิน {fmt_num(rows[-1]['sessions_per_day'])} คัน/วันต่อเนื่อง ค่อยพิจารณาขยายเฟสถัดไป",
        "สำหรับการคุยกับเจ้าของพื้นที่ เอกสารควรเน้นว่า GP เป็นรายได้เสริมระยะยาว โดยไม่ต้องรับความเสี่ยงด้าน Capex และการเดินระบบ",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def add_caveats(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อควรใช้ในการตัดสินใจ")

    bullets = [
        "เอกสารนี้ใช้ Base Case และใช้การเติบโตตาม area-demand growth ของ TH-EVI เพื่อให้เห็นภาพ 10 ปีที่ไม่แบนเกินไป",
        "โมเดล point estimate เดิมของจุดนี้อยู่ที่ประมาณ 6 คัน/วันในปีแรก และจัดเป็น low relevance ดังนั้นเอกสารนี้ควรใช้เพื่อคุยเบื้องต้น ไม่ใช่ยืนยันลงทุนทันที",
        "โมเดลลงทุนเองรวมสมมติฐานค่าไฟ 4 บาท/kWh, CPO GP 8%, และ O&M 36,000 บาท/ปี แล้ว",
        "โมเดล GP 0.25 บาท/kWh ถือว่าเจ้าของพื้นที่ไม่ลงทุน Capex และรับรายได้ตามพลังงานที่ขายได้จริง",
        "หากต้องการความแม่นยำเพิ่ม ควรมีการเก็บทราฟฟิกจริง, ลักษณะที่จอด, และศักยภาพการเข้าออกของพื้นที่ก่อนตัดสินใจ",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    rows, base_result, payback = projection_rows()
    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    add_executive_summary(doc, rows[0], rows[-1], base_result, payback)
    add_key_assumptions(doc)
    add_area_analysis(doc, base_result)
    add_comparison_table(doc, rows)
    add_recommendation(doc, rows, payback)
    add_caveats(doc)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
