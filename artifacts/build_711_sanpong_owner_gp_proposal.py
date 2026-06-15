from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from th_evi.spatial import analyze_click_location


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\7-11_Sanpong_Owner_GP_Proposal.docx")

SITE_NAME = "7-11 สันโป่ง"
ADDRESS = "San Pong, Mae Rim District, Chiang Mai 50180"
LAT = 18.942782
LON = 98.942951
PROVINCE = "Chiang Mai"
START_YEAR = 2026
END_YEAR = 2035
SELL_PRICE_PER_KWH = 7.9
OWNER_GP_PER_KWH = 0.25
AVG_KWH_PER_CAR = 35.0
RECOMMENDED_POWER_KW = 120
RECOMMENDED_CONNECTORS = 2

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)
GREEN = RGBColor(27, 94, 32)
AMBER = RGBColor(180, 83, 9)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def fmt_num(value: float, digits: int = 1) -> str:
    return f"{value:,.{digits}f}"


def fmt_int(value: float) -> str:
    return f"{round(value):,}"


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 12.5, ACCENT),
        ("Heading 3", 11.5, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | Owner GP Proposal")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(SITE_NAME)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def annual_projection_rows() -> tuple[list[dict], dict]:
    rows: list[dict] = []
    cumulative_gp = 0.0
    first_result: dict | None = None
    for year in range(START_YEAR, END_YEAR + 1):
        result = analyze_click_location(
            lat=LAT,
            lon=LON,
            province=PROVINCE,
            year=year,
            scenario="base",
            mode="urban",
            avg_kwh_per_session=AVG_KWH_PER_CAR,
            price_per_kwh=SELL_PRICE_PER_KWH,
        )
        if first_result is None:
            first_result = result
        cars_per_day = float(result["net_sessions_per_day"])
        daily_kwh = cars_per_day * AVG_KWH_PER_CAR
        annual_kwh = daily_kwh * 365
        annual_revenue = annual_kwh * SELL_PRICE_PER_KWH
        annual_owner_gp = annual_kwh * OWNER_GP_PER_KWH
        cumulative_gp += annual_owner_gp
        rows.append({
            "year": year,
            "cars_per_day": cars_per_day,
            "daily_kwh": daily_kwh,
            "annual_kwh": annual_kwh,
            "annual_revenue": annual_revenue,
            "annual_owner_gp": annual_owner_gp,
            "cumulative_owner_gp": cumulative_gp,
            "utilization_pct": (daily_kwh / (RECOMMENDED_POWER_KW * 24.0)) * 100.0,
        })
    return rows, first_result or {}


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis + GP Proposal")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("7-11 สันโป่ง | ข้อเสนอ GP สำหรับเจ้าของพื้นที่")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.7, 4.7])
    rows = [
        ("สถานที่", SITE_NAME),
        ("ที่อยู่", ADDRESS),
        ("พิกัด", f"{LAT:.6f}, {LON:.6f}"),
        ("รูปแบบความร่วมมือ", "เจ้าของพื้นที่ไม่ลงทุน | รับ GP 0.25 บาท/kWh"),
        ("ข้อเสนอสถานี", f"SINEXCEL {RECOMMENDED_POWER_KW} kW | {RECOMMENDED_CONNECTORS} หัวชาร์จ"),
        ("ช่วงพยากรณ์", f"{START_YEAR}-{END_YEAR} รวม 10 ปี"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_executive_summary(doc: Document, first_row: dict, final_row: dict, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    text = (
        f"จุด 7-11 สันโป่งอยู่ในแนวแม่ริม-โชตนา-Route 107 ซึ่งมีทราฟฟิกผ่านจริงและมีดีมานด์เชิงคอร์ริดอร์อยู่ระดับหนึ่ง "
        f"แต่ในเชิงคุณภาพ โมเดลยังจัดพื้นที่นี้เป็น low relevance เพราะสัญญาณ access และ activity ติดจุดยังไม่แข็งมากพอเมื่อเทียบกับโหนดเมืองชัด ๆ "
        f"ดังนั้นโอกาสของเจ้าของพื้นที่จึงเหมาะกับโมเดล “รับ GP โดยไม่ลงทุนเอง” มากกว่าการคาดหวังรายได้ก้อนใหญ่จากไซต์ระดับ flagship "
        f"ใน Base Case ปี {START_YEAR} คาดว่ามีรถ EV เข้ามาชาร์จประมาณ {fmt_num(first_row['cars_per_day'])} คัน/วัน "
        f"คิดเป็นรายได้รวมของสถานีประมาณ {fmt_int(first_row['annual_revenue'])} บาท/ปี และทำให้เจ้าของพื้นที่ได้รับ GP ประมาณ {fmt_int(first_row['annual_owner_gp'])} บาท/ปี "
        f"ตลอดช่วง 10 ปีคาดว่า GP สะสมจะอยู่ที่ประมาณ {fmt_int(final_row['cumulative_owner_gp'])} บาท "
        f"ข้อเสนอที่เหมาะสมกับไซต์นี้คือเริ่มด้วยสถานี SINEXCEL {RECOMMENDED_POWER_KW} kW จำนวน {RECOMMENDED_CONNECTORS} หัวชาร์จ "
        f"เพราะยังพอให้ภาพ fast-charging ที่ลูกค้ายอมรับได้ แต่ไม่ใหญ่เกินดีมานด์ของจุดในปัจจุบัน"
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def add_key_assumptions(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สมมติฐานหลักของข้อเสนอ")

    table = doc.add_table(rows=8, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [1.9, 1.4, 2.7])
    headers = ["รายการ", "ค่า", "หมายเหตุ"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("รูปแบบรายได้", "GP 0.25 บาท/kWh", "เจ้าของพื้นที่รับตามพลังงานที่ขายได้จริง"),
        ("เงินลงทุนฝั่งเจ้าของ", "0 บาท", "สมมติว่าผู้ลงทุนเป็นผู้รับ CAPEX ทั้งหมด"),
        ("กรอบเวลา", "10 ปี", "ใช้ดูรายได้ GP ระยะยาวของเจ้าของไซต์"),
        ("สถานีที่แนะนำ", f"{RECOMMENDED_POWER_KW} kW / {RECOMMENDED_CONNECTORS} หัว", "ขนาดเหมาะสมกับ demand และ perception ของลูกค้าหน้าร้านสะดวกซื้อ"),
        ("ดีมานด์ต่อคัน", "35 kWh/คัน", "ใช้เป็นพลังงานเฉลี่ยต่อการชาร์จ 1 คัน"),
        ("ราคาขาย", "7.9 บาท/kWh", "ใช้เพื่อแสดงรายได้รวมของสถานีประกอบการนำเสนอ"),
        ("สถานีเดิมในไซต์", "ไม่มีการยืนยันในระบบ", "ข้อเสนอนี้อิงจากพิกัด 7-11 สันโป่งโดยตรง"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK


def add_area_analysis(doc: Document, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Area Analysis")

    doc.add_paragraph(
        "จุด 7-11 สันโป่งอยู่ในแนวแม่ริมซึ่งได้แรงจากคอร์ริดอร์ Mae Rim / Chotana / Route 107 และ node ของอำเภอแม่ริม "
        "แต่ ณ พิกัดนี้ยังไม่ได้อยู่ติดแหล่งกิจกรรมเมืองหนาแน่นหรือ anchor ขนาดใหญ่ในระยะใกล้มากพอ ทำให้โมเดลให้ผลเป็น low relevance "
        "ความหมายในเชิงธุรกิจคือมีรถผ่านและมีโอกาสเกิดการชาร์จจริง แต่ยังไม่ใช่ไซต์ที่ควรเริ่มด้วยสถานีขนาดใหญ่เกินจำเป็น"
    )

    metrics = doc.add_table(rows=8, cols=3)
    metrics.style = "Table Grid"
    metrics.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(metrics, [2.0, 1.2, 2.8])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมาย"]
    for col, text in enumerate(headers):
        cell = metrics.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะพื้นที่", str(base_result.get("eligibility_status", "-")).title(), str(base_result.get("eligibility_reason", "-"))),
        ("ประเภททำเล", str(base_result.get("location_type", "-")).title(), "โมเดลอ่านเป็น highway-side opportunity มากกว่า urban node"),
        ("ดีมานด์รวมของพื้นที่", f"{fmt_num(base_result.get('gross_area_demand_sessions', 0))} คัน/วัน", "สะท้อน demand pool ของพื้นที่แม่ริมรอบกว้าง"),
        ("ดีมานด์ของจุด", f"{fmt_num(base_result.get('net_sessions_per_day', 0))} คัน/วัน", "จำนวนรถที่คาดว่าจุดนี้จะดึงเข้าชาร์จได้จริงใน Base Case"),
        ("พลังงานต่อวัน", f"{fmt_num(base_result.get('daily_kwh', 0))} kWh/วัน", f"คำนวณจาก {AVG_KWH_PER_CAR:.0f} kWh/คัน"),
        ("รายได้รวมของสถานีต่อวัน", f"{fmt_int(base_result.get('daily_revenue', 0))} บาท/วัน", "คำนวณที่ราคาขาย 7.9 บาท/kWh"),
        ("GP เจ้าของพื้นที่ต่อวัน", f"{fmt_int(base_result.get('daily_kwh', 0) * OWNER_GP_PER_KWH)} บาท/วัน", "คำนวณที่ GP 0.25 บาท/kWh"),
    ]
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_values):
            cell = metrics.cell(row_idx, col_idx)
            cell.text = text
            set_cell_border(cell)
            if col_idx == 0:
                set_cell_shading(cell, "F8FAFC")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = INK

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("หลักฐานสนับสนุนและข้อสังเกต")

    bullets = [
        f"Hot zone เด่น: {', '.join(item['name'] for item in base_result.get('top_zones', [])[:3]) or 'ไม่มี'}",
        f"Business area เด่น: {', '.join(item['name'] for item in base_result.get('top_business_areas', [])[:2]) or 'ไม่มี'}",
        f"คู่แข่งใกล้สุดในระบบห่างประมาณ {fmt_num(base_result.get('nearest_competitor_km', 0))} กม.",
        f"Access anchor ใกล้สุดห่างประมาณ {fmt_num(base_result.get('nearest_access_anchor_km', 0))} กม.",
        "ข้อสรุปเชิงกลยุทธ์: จุดนี้เหมาะกับข้อเสนอ GP แบบเบาเงินลงทุนของเจ้าของไซต์ แต่ยังไม่ใช่ฐานที่ควรผลักสถานีขนาดใหญ่มากตั้งแต่วันแรก",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def add_financial_forecast(doc: Document, rows: list[dict]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Base Case Forecast และรายได้ GP ของเจ้าของพื้นที่")

    note = doc.add_paragraph()
    note.add_run(
        "สูตรที่ใช้: จำนวนรถชาร์จต่อวัน = sessions/day จาก TH-EVI | พลังงานต่อวัน = รถชาร์จต่อวัน x 35 kWh | "
        "รายได้รวมของสถานี = พลังงาน x 7.9 บาท/kWh | GP ของเจ้าของพื้นที่ = พลังงาน x 0.25 บาท/kWh"
    )

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [0.55, 0.85, 0.95, 1.05, 1.15, 1.15, 0.9])
    headers = ["ปี", "คัน/วัน", "kWh/วัน", "รายได้/ปี", "GP/ปี", "GP สะสม", "Utilization"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    for row in rows:
        values = [
            str(row["year"]),
            fmt_num(row["cars_per_day"]),
            fmt_num(row["daily_kwh"]),
            fmt_int(row["annual_revenue"]),
            fmt_int(row["annual_owner_gp"]),
            fmt_int(row["cumulative_owner_gp"]),
            f"{fmt_num(row['utilization_pct'])}%",
        ]
        row_cells = table.add_row().cells
        for idx, value in enumerate(values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("ข้อสรุปด้านรายได้ของเจ้าของพื้นที่")

    first_row = rows[0]
    final_row = rows[-1]
    summary_rows = [
        f"ปีแรก ({START_YEAR}) คาดว่ามีรถ EV เข้ามาชาร์จประมาณ {fmt_num(first_row['cars_per_day'])} คัน/วัน ทำให้ Wasabi Park ได้ GP ประมาณ {fmt_int(first_row['annual_owner_gp'])} บาท/ปี",
        f"ปีสุดท้ายของช่วงพยากรณ์ ({END_YEAR}) คาดว่ามีรถ EV เข้ามาชาร์จประมาณ {fmt_num(final_row['cars_per_day'])} คัน/วัน และทำให้ Wasabi Park ได้ GP ประมาณ {fmt_int(final_row['annual_owner_gp'])} บาท/ปี",
        f"GP สะสมตลอด 10 ปีประมาณ {fmt_int(final_row['cumulative_owner_gp'])} บาท",
        f"เมื่อเทียบกับศักยภาพของไซต์ ข้อเสนอ GP แบบไม่ต้องลงทุนเองถือว่าเหมาะกว่าโมเดลที่เจ้าของพื้นที่ลงเงินก้อนใหญ่ด้วยตัวเอง",
    ]
    for item in summary_rows:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        if "GP สะสม" in item:
            run.font.color.rgb = GREEN
            run.bold = True


def add_recommendation(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("คำแนะนำขนาดสถานี")

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = callout.cell(0, 0)
    set_cell_border(cell, color="E8B26A", size="10")
    set_cell_shading(cell, "FFF7ED")
    p = cell.paragraphs[0]
    run = p.add_run(
        f"คำแนะนำ: เหมาะเริ่มที่ SINEXCEL {RECOMMENDED_POWER_KW} kW จำนวน {RECOMMENDED_CONNECTORS} หัวชาร์จ "
        f"ไม่แนะนำเริ่มด้วย 180 kW 4 ช่องจอดสำหรับจุดนี้ เพราะดีมานด์ของจุดยังอยู่เพียงประมาณ 6 คัน/วันใน Base Case "
        f"สถานี 120 kW ยังพอให้ภาพ fast charging ที่ลูกค้ารับรู้ได้ ขณะเดียวกันไม่ทำให้โครงการใหญ่เกินความต้องการของไซต์"
    )
    run.font.color.rgb = AMBER
    run.bold = True

    bullets = [
        "ถ้าต้องการเป้าหมายหลักคือสร้างรายได้ GP ให้เจ้าของพื้นที่ ควรเริ่มขนาดพอดีและพิสูจน์ demand ก่อน",
        "ถ้า demand ในอนาคตสูงกว่าที่คาดและมีการใช้งานเกิน 10-12 คัน/วันต่อเนื่อง ค่อยพิจารณาขยายจำนวนหัวหรือยกระดับกำลังไฟในเฟสถัดไป",
        "ข้อเสนอ 180 kW 4 ช่องจอดจะเหมาะกว่าก็ต่อเมื่อมีหลักฐานทราฟฟิก EV หน้างานจริงเพิ่มขึ้นชัดเจน หรือมีแผนเชิงเครือข่ายจากผู้ลงทุน",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def add_caveats(doc: Document, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อควรใช้ในการตัดสินใจ")

    bullets = [
        "เอกสารฉบับนี้ใช้ Base Case เพียงกรณีเดียว ตามที่ร้องขอ และใช้ราคาขาย 7.9 บาท/kWh เพื่อแสดงรายได้รวมของสถานีประกอบการนำเสนอ",
        "รายได้ GP ของเจ้าของพื้นที่ในเอกสารนี้คำนวณจาก 0.25 บาท/kWh ของพลังงานที่ขายได้จริง โดยไม่ได้สมมติให้เจ้าของต้องรับภาระค่าไฟหรือค่าใช้จ่ายระบบ",
        "ผลของโมเดลชี้ว่าพื้นที่นี้มี demand บางส่วนจาก corridor และ node ของแม่ริม แต่ direct access signal ยังอ่อน จึงควรนำเสนอเจ้าของแบบระมัดระวังและไม่คาดหวังสถานีใหญ่ตั้งแต่เริ่ม",
        f"ระดับความเชื่อมั่นของจุดนี้ในระบบอยู่ที่ {base_result.get('confidence', '-')}",
        "ถ้าต้องการความแม่นยำสูงขึ้น ควรลงพื้นที่ตรวจจำนวนรถเข้าออกจริง ลักษณะที่จอด และความสะดวกในการเข้าใช้หน้าร้านก่อนปิดดีล",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    rows, base_result = annual_projection_rows()
    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    add_executive_summary(doc, rows[0], rows[-1], base_result)
    add_key_assumptions(doc)
    add_area_analysis(doc, base_result)
    add_financial_forecast(doc, rows)
    add_recommendation(doc)
    add_caveats(doc, base_result)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
