from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from th_evi.spatial import analyze_click_location


OUT_PATH = Path(r"D:\Work\TH-EVI\artifacts\Ratchapruek_Market_Owner_Analysis.docx")

SITE_NAME = "ตลาดนัดราชพฤกษ์"
ADDRESS = "R79C+M6P, Tambon Nai Mueang, Mueang Phitsanulok District, Phitsanulok 65000"
LAT = 16.819181
LON = 100.268779
PROVINCE = "Phitsanulok"
START_YEAR = 2026
END_YEAR = 2035
SELL_PRICE_PER_KWH = 7.9
ELECTRICITY_COST_PER_KWH = 4.0
AVG_KWH_PER_CAR = 35.0
CPO_GP_RATE = 0.08
O_AND_M_PER_YEAR = 36_000

CONFIDENT_CASE_CARS_PER_DAY_BY_YEAR = {
    2026: 22.0,
    2027: 24.0,
    2028: 26.0,
    2029: 28.0,
    2030: 30.0,
    2031: 32.0,
    2032: 34.0,
    2033: 35.0,
    2034: 36.0,
    2035: 38.0,
}

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)
GREEN = RGBColor(27, 94, 32)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def fmt_num(value: float, digits: int = 1) -> str:
    return f"{value:,.{digits}f}"


def fmt_int(value: float) -> str:
    return f"{round(value):,}"


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 12.5, ACCENT),
        ("Heading 3", 11.5, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | Area Analysis")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("ตลาดนัดราชพฤกษ์")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def annual_projection_rows() -> list[dict]:
    rows: list[dict] = []
    cumulative_revenue = 0.0
    cumulative_operating_cf = 0.0
    for year in range(START_YEAR, END_YEAR + 1):
        result = analyze_click_location(
            lat=LAT,
            lon=LON,
            province=PROVINCE,
            year=year,
            scenario="base",
            mode="urban",
            avg_kwh_per_session=AVG_KWH_PER_CAR,
            price_per_kwh=SELL_PRICE_PER_KWH,
        )
        modeled_cars_per_day = float(result["net_sessions_per_day"])
        cars_per_day = min(modeled_cars_per_day, CONFIDENT_CASE_CARS_PER_DAY_BY_YEAR[year])
        daily_kwh = cars_per_day * AVG_KWH_PER_CAR
        annual_kwh = daily_kwh * 365.0
        annual_revenue = annual_kwh * SELL_PRICE_PER_KWH
        annual_cpo_gp = annual_revenue * CPO_GP_RATE
        annual_electricity_cost = annual_kwh * ELECTRICITY_COST_PER_KWH
        annual_operating_cf = annual_revenue - annual_cpo_gp - annual_electricity_cost - O_AND_M_PER_YEAR
        cumulative_revenue += annual_revenue
        cumulative_operating_cf += annual_operating_cf
        rows.append(
            {
                "year": year,
                "modeled_cars_per_day": modeled_cars_per_day,
                "cars_per_day": cars_per_day,
                "daily_kwh": daily_kwh,
                "annual_kwh": annual_kwh,
                "annual_revenue": annual_revenue,
                "annual_cpo_gp": annual_cpo_gp,
                "annual_electricity_cost": annual_electricity_cost,
                "annual_operating_cf": annual_operating_cf,
                "cumulative_revenue": cumulative_revenue,
                "cumulative_operating_cf": cumulative_operating_cf,
                "utilization_240_pct": (daily_kwh / (240.0 * 24.0)) * 100.0,
                "utilization_360_pct": (daily_kwh / (360.0 * 24.0)) * 100.0,
            }
        )
    return rows


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis สำหรับเจ้าของสถานที่")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("ตลาดนัดราชพฤกษ์ | เปรียบเทียบแนวคิด 120 kW กับ 180 kW และข้อเสนอแนะขนาดสถานีแบบ Confident Case")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    set_repeat_table_widths(meta, [1.7, 4.7])
    rows = [
        ("สถานที่", SITE_NAME),
        ("ที่อยู่", ADDRESS),
        ("พิกัด", f"{LAT:.6f}, {LON:.6f}"),
        ("มุมมองการวิเคราะห์", "เจ้าของสถานที่ลงทุนเอง"),
        ("ตัวเลือกที่พิจารณา", "DC 120 kW จำนวน 2 ตู้ 4 ช่องจอด หรือ DC 180 kW จำนวน 2 ตู้ 4 ช่องจอด"),
        ("ช่วงพยากรณ์", f"{START_YEAR}-{END_YEAR} รวม 10 ปี"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
        set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def add_executive_summary(doc: Document, base_result: dict, first_row: dict, final_row: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    set_cell_border(cell, color="C9D6E3", size="10")
    set_cell_shading(cell, "F4F7FA")
    text = (
        f"ตลาดนัดราชพฤกษ์อยู่ในแกนเมืองพิษณุโลกที่ค่อนข้างแรง ทั้งจากตลาด ไลฟ์สไตล์ แกนเมืองเก่า สถานีรถไฟ "
        f"และศูนย์ค้าปลีกใกล้เคียง โดยโมเดลอ่านจุดนี้เป็น {base_result['location_type']} และผ่านเกณฑ์ "
        f"{base_result['eligibility_status']} ชัดเจน ในปี {START_YEAR} โมเดลให้ดีมานด์ที่จุดประมาณ "
        f"{fmt_num(base_result['net_sessions_per_day'])} คัน/วัน แต่เอกสารฉบับนี้ใช้ Confident Case ที่อยู่ระหว่าง Base กับ Best และยังค่อนมาทาง Base เพียง "
        f"{fmt_num(first_row['cars_per_day'])} คัน/วันในปีแรก เพื่อเผื่อการแข่งขัน การ ramp-up ของลูกค้า และการใช้งานจริงของสถานี "
        f"4 ช่องจอด ข้อเสนอแนะของเราคือให้เริ่มที่ DC 180 kW จำนวน 2 ตู้ 4 ช่องจอด มากกว่า 120 kW จำนวน 2 ตู้ 4 ช่องจอด "
        f"เพราะทำเลนี้เป็นเมืองชั้นในที่ลูกค้าให้ความสำคัญกับความมั่นใจและความเร็วในการชาร์จ การใช้ 180 kW จะช่วยเรื่องภาพลักษณ์ ลด queue anxiety "
        f"และรองรับช่วงพีคได้ดีกว่า โดยยังอยู่ในกรอบคาดการณ์ที่ระมัดระวังพอ ตลอด 10 ปีนี้ หากทำยอดตาม Confident Case "
        f"รายได้ขั้นต้นสะสมจะอยู่ที่ประมาณ {fmt_int(final_row['cumulative_revenue'])} บาท และกระแสเงินสดจากการดำเนินงานสะสมหลังหักค่าไฟ, CPO GP และ O&M "
        f"จะอยู่ที่ประมาณ {fmt_int(final_row['cumulative_operating_cf'])} บาท"
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def add_key_assumptions(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สมมติฐานหลักของ Confident Case")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [2.0, 1.3, 3.0])
    headers = ["รายการ", "ค่า", "หมายเหตุ"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("ราคาขาย", "7.9 บาท/kWh", "ใช้เรทราคาขายปกติ"),
        ("ต้นทุนค่าไฟ", "4.0 บาท/kWh", "สมมติฐานต้นทุนพลังงาน"),
        ("พลังงานเฉลี่ยต่อคัน", "35 kWh/คัน", "ใช้เป็นฐานแปลงจากคัน/วันเป็นรายได้"),
        ("CPO GP", "8% ของรายได้", "หักออกก่อนคำนวณกระแสเงินสดจากการดำเนินงาน"),
        ("O&M", "36,000 บาท/ปี", "คิดเทียบ 3,000 บาท/เดือน"),
        ("กรอบการพยากรณ์", "10 ปี", "ตั้งแต่ 2026 ถึง 2035"),
        ("Confident Case คัน/วัน", "22 -> 38 คัน/วัน", "ค่อย ๆ เติบโตตามปี และยังต่ำกว่าค่าดิบจากโมเดลเพื่อให้ไม่ aggressive เกินไป"),
        ("ตัวเลือกสถานี", "120 kW หรือ 180 kW", "ทั้งสองแบบใช้ 2 ตู้ 4 ช่องจอด"),
        ("ข้อเสนอแนะ", "180 kW", "เน้น perception, turnover และการแข่งขันในโซนเมือง"),
    ]
    for row_values in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK


def add_area_analysis(doc: Document, base_result: dict, first_row: dict, final_row: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Area Analysis")

    doc.add_paragraph(
        "พื้นที่นี้อยู่ในแกนเมืองพิษณุโลกที่มีความหนาแน่นของกิจกรรมเชิงพาณิชย์และการเดินทางสูงกว่าค่าเฉลี่ยของจังหวัด "
        "โดยจุดนี้ได้รับแรงหนุนจากตลาด-เมืองเก่า-รถไฟ-ค้าปลีก และยังอยู่ใกล้คู่แข่งที่เป็นสถานีชาร์จจริงหลายจุด "
        "ทำให้เป็นไซต์ที่เหมาะกับการมองแบบ destination / urban charging มากกว่าการมองเป็นปั๊มทางผ่าน"
    )

    metrics = doc.add_table(rows=1, cols=3)
    metrics.style = "Table Grid"
    metrics.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(metrics, [2.1, 1.2, 2.9])
    headers = ["ตัวชี้วัด", "ค่า", "ความหมาย"]
    for col, text in enumerate(headers):
        cell = metrics.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        ("สถานะพื้นที่", str(base_result["eligibility_status"]).title(), base_result["eligibility_reason"]),
        ("ประเภททำเล", str(base_result["location_type"]).title(), "อ่านเป็นจุดหมายปลายทางในเมือง ไม่ใช่ highway pass-through"),
        ("ดีมานด์รวมของพื้นที่", f"{fmt_num(base_result['gross_area_demand_sessions'])} คัน/วัน", "สะท้อน demand pool รอบจุดในภาพรวม"),
        ("ดีมานด์ดิบของโมเดลที่จุด", f"{fmt_num(base_result['net_sessions_per_day'])} คัน/วัน", "ค่าดิบของโมเดลก่อน cap ให้ระมัดระวัง"),
        ("Confident Case ปีแรก", f"{fmt_num(first_row['cars_per_day'])} คัน/วัน", "ใช้แทน forecast ปีแรกของเอกสารฉบับนี้"),
        ("Confident Case ปี 2030", f"{fmt_num([r for r in annual_projection_rows() if r['year'] == 2030][0]['cars_per_day'])} คัน/วัน", "ปีที่ demand เมืองพิษณุโลกเริ่มหนาแน่นขึ้นชัด"),
        ("Confident Case ปีสุดท้าย", f"{fmt_num(final_row['cars_per_day'])} คัน/วัน", "ใช้ดูความพอดีของขนาดสถานีระยะ 10 ปี"),
        ("ระยะใกล้คู่แข่ง", f"{fmt_num(base_result['nearest_competitor_km'])} กม.", "มี competition จริงในเมือง จึงควรใช้สเปกที่ลูกค้ารู้สึกว่าเร็วพอ"),
    ]
    for row_values in rows:
        row_cells = metrics.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("แรงหนุนหลักของพื้นที่")

    bullets = [
        f"POI เด่น: {', '.join(item['name'] for item in base_result['top_pois'][:5])}",
        f"Hot zone เด่น: {', '.join(item['name'] for item in base_result['top_zones'][:3])}",
        f"Business area เด่น: {', '.join(item['name'] for item in base_result['top_business_areas'][:2])}",
        f"คู่แข่งหลัก: {', '.join(item['name'] for item in base_result['top_competitors'][:5])}",
        f"Nearest access anchor {fmt_num(base_result['nearest_access_anchor_km'])} กม. | nearest competitor {fmt_num(base_result['nearest_competitor_km'])} กม.",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def add_sizing_recommendation(doc: Document, first_row: dict, final_row: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อเสนอแนะขนาดสถานี")

    doc.add_paragraph(
        "ทั้ง 120 kW และ 180 kW จำนวน 2 ตู้ 4 ช่องจอด สามารถทำงานได้ในเชิงเทคนิคสำหรับ Confident Case นี้ "
        "แต่หากมองจากการแข่งขันในเมือง ความคาดหวังของลูกค้า และความต่อเนื่องของ demand ใน 10 ปี "
        "เรามองว่า 180 kW จำนวน 2 ตู้ 4 ช่องจอดเหมาะกว่า"
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [2.2, 1.5, 1.5, 2.0])
    headers = ["หัวข้อ", "120 kW x 2 ตู้", "180 kW x 2 ตู้", "ความเห็น"]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    rows = [
        (
            "ภาพลักษณ์ต่อผู้ใช้",
            "พอใช้ได้",
            "ดีกว่า",
            "180 kW ช่วยเรื่อง perception ว่าเร็วและน่าแวะกว่า",
        ),
        (
            "ความเสี่ยงคิวช่วงพีค",
            "ตึงกว่า",
            "ผ่อนกว่า",
            "แม้คัน/วันใน Confident Case ยังไม่สูงสุด แต่ช่วงพีคของตลาด/เมืองอาจมากระจุกได้",
        ),
        (
            "Utilization ปีแรก",
            f"{fmt_num(first_row['utilization_240_pct'])}%",
            f"{fmt_num(first_row['utilization_360_pct'])}%",
            "ทั้งสองแบบยังไม่แน่นเกินไป แต่ 180 kW สบายกว่า",
        ),
        (
            "Utilization ปีสุดท้าย",
            f"{fmt_num(final_row['utilization_240_pct'])}%",
            f"{fmt_num(final_row['utilization_360_pct'])}%",
            "ปีท้าย ๆ 120 kW ยังทำได้ แต่ 180 kW ให้ margin และภาพลักษณ์ที่ดีกว่า",
        ),
        (
            "ข้อเสนอแนะ",
            "สำรอง",
            "แนะนำ",
            "เลือก 180 kW x 2 ตู้ 4 ช่องจอด",
        ),
    ]
    for row_values in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK
        if row_values[0] == "ข้อเสนอแนะ":
            row_cells[2].paragraphs[0].runs[0].font.bold = True
            row_cells[2].paragraphs[0].runs[0].font.color.rgb = GREEN


def add_financial_forecast(doc: Document, rows: list[dict]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Confident Case Forecast 10 ปี")

    doc.add_paragraph(
        "สูตรที่ใช้: พลังงานต่อวัน = คัน/วัน x 35 kWh | รายได้ขั้นต้น = พลังงาน x 7.9 บาท/kWh | "
        "CPO GP = 8% ของรายได้ | ค่าไฟ = พลังงาน x 4.0 บาท/kWh | O&M = 36,000 บาท/ปี | "
        "กระแสเงินสดจากการดำเนินงาน = รายได้ขั้นต้น - CPO GP - ค่าไฟ - O&M"
    )

    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_widths(table, [0.55, 0.72, 0.82, 0.95, 0.95, 0.9, 0.95, 1.0, 1.0])
    headers = [
        "ปี",
        "คัน/วัน",
        "kWh/วัน",
        "รายได้/ปี",
        "ค่าไฟ/ปี",
        "CPO GP",
        "CF/ปี",
        "รายได้สะสม",
        "CF สะสม",
    ]
    for col, text in enumerate(headers):
        cell = table.cell(0, col)
        cell.text = text
        set_cell_border(cell)
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = INK

    for row in rows:
        values = [
            str(row["year"]),
            fmt_num(row["cars_per_day"]),
            fmt_num(row["daily_kwh"]),
            fmt_int(row["annual_revenue"]),
            fmt_int(row["annual_electricity_cost"]),
            fmt_int(row["annual_cpo_gp"]),
            fmt_int(row["annual_operating_cf"]),
            fmt_int(row["cumulative_revenue"]),
            fmt_int(row["cumulative_operating_cf"]),
        ]
        row_cells = table.add_row().cells
        for idx, value in enumerate(values):
            row_cells[idx].text = value
            set_cell_border(row_cells[idx])
            if idx == 0:
                set_cell_shading(row_cells[idx], "F8FAFC")
                row_cells[idx].paragraphs[0].runs[0].font.bold = True
                row_cells[idx].paragraphs[0].runs[0].font.color.rgb = INK

    first_row = rows[0]
    final_row = rows[-1]
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run("ข้อสรุปทางการเงิน")

    bullets = [
        f"ปีแรก ({START_YEAR}) Confident Case ใช้ {fmt_num(first_row['cars_per_day'])} คัน/วัน คิดเป็นรายได้ขั้นต้นประมาณ {fmt_int(first_row['annual_revenue'])} บาท/ปี",
        f"ปี 2030 Confident Case ใช้ {fmt_num([r for r in rows if r['year'] == 2030][0]['cars_per_day'])} คัน/วัน ซึ่งเท่ากับประมาณ {fmt_int([r for r in rows if r['year'] == 2030][0]['annual_revenue'])} บาท/ปี",
        f"ปีสุดท้าย ({END_YEAR}) Confident Case ใช้ {fmt_num(final_row['cars_per_day'])} คัน/วัน คิดเป็นรายได้ขั้นต้นประมาณ {fmt_int(final_row['annual_revenue'])} บาท/ปี",
        f"รายได้ขั้นต้นสะสม 10 ปีประมาณ {fmt_int(final_row['cumulative_revenue'])} บาท",
        f"กระแสเงินสดจากการดำเนินงานสะสม 10 ปีประมาณ {fmt_int(final_row['cumulative_operating_cf'])} บาท",
    ]
    for item in bullets:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        if "สะสม 10 ปี" in item:
            run.font.color.rgb = GREEN
            run.bold = True


def add_caveats(doc: Document, base_result: dict) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อควรใช้ในการตัดสินใจ")

    bullets = [
        "เอกสารฉบับนี้ใช้ Confident Case ซึ่งอยู่ระหว่าง Base กับ Best แต่ยังค่อนมาทาง Base เพื่อให้ตัวเลขดูมีแรงพอสำหรับคุยเจ้าของพื้นที่และยังอธิบายได้",
        "รายได้และกระแสเงินสดในเอกสารนี้ยังไม่รวม CAPEX เพราะโจทย์รอบนี้ต้องการดูความเหมาะของสเปกและรายได้จากการใช้งานเป็นหลัก",
        "หากเจ้าของต้องการตัดสินใจลงทุนจริง ควรเติมราคาตู้, งานไฟ, หม้อแปลง, งานโยธา และค่าเชื่อมระบบ เพื่อทำ payback และ IRR ต่อ",
        "เหตุผลที่แนะนำ 180 kW ไม่ได้มาจาก average utilization อย่างเดียว แต่มาจาก perception ของลูกค้า, การแข่งขันในเมือง, และการลดความกังวลเรื่องคิว",
        f"ระดับความเชื่อมั่นของจุดนี้ในระบบอยู่ที่ {base_result['confidence']} และคำเตือนจากระบบคือ: {', '.join(base_result['warnings']) if base_result['warnings'] else 'ไม่มี'}",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    base_result = analyze_click_location(
        lat=LAT,
        lon=LON,
        province=PROVINCE,
        year=START_YEAR,
        scenario="base",
        mode="urban",
        avg_kwh_per_session=AVG_KWH_PER_CAR,
        price_per_kwh=SELL_PRICE_PER_KWH,
    )
    rows = annual_projection_rows()
    first_row = rows[0]
    final_row = rows[-1]

    doc = Document()
    style_doc(doc)
    add_title_block(doc)
    add_executive_summary(doc, base_result, first_row, final_row)
    add_key_assumptions(doc)
    add_area_analysis(doc, base_result, first_row, final_row)
    add_sizing_recommendation(doc, first_row, final_row)
    add_financial_forecast(doc, rows)
    add_caveats(doc, base_result)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
