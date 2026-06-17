"""Word report generator for owner-facing area analysis from heatmap/click points."""

from __future__ import annotations

from dataclasses import asdict, dataclass
from datetime import datetime
import html
import importlib.util
import io
import json
import math
import os
from pathlib import Path
import subprocess
import sys
import tempfile
from typing import Any
import re
import struct
import urllib.request
import zlib

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
if os.environ.get("VERCEL"):
    DEFAULT_REPORT_OUTPUT_DIR = Path(tempfile.gettempdir()) / "th-evi-generated-reports"
else:
    DEFAULT_REPORT_OUTPUT_DIR = REPO_ROOT / "artifacts" / "generated_reports"
REPORT_OUTPUT_DIR = Path(os.environ.get("TH_EVI_REPORT_OUTPUT_DIR", str(DEFAULT_REPORT_OUTPUT_DIR)))

BLACK = RGBColor(0, 0, 0)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 99, 110)
ACCENT = RGBColor(46, 116, 181)
GREEN = RGBColor(27, 94, 32)

MAP_BG = (245, 248, 252)
MAP_GRID = (225, 232, 240)
MAP_BORDER = (180, 194, 208)
MAP_POI = (37, 99, 235)
MAP_COMPETITOR = (124, 58, 237)
MAP_ZONE = (245, 158, 11)
MAP_BUSINESS = (14, 165, 233)
MAP_SITE = (239, 68, 68)
HEAT_COLORS = [
    (0.30, (255, 243, 191)),
    (0.45, (255, 209, 102)),
    (0.60, (248, 150, 30)),
    (0.75, (239, 71, 111)),
    (0.90, (181, 23, 158)),
    (1.00, (123, 44, 191)),
]
OSM_TILE_TEMPLATE = "https://tile.openstreetmap.org/{z}/{x}/{y}.png"
OSM_USER_AGENT = "TH-EVI/1.0 (planning-report; contact pongwara.ja@tceproject.co.th)"


@dataclass(slots=True)
class OwnerAreaReportRequest:
    site_name: str | None
    province: str
    lat: float
    lon: float
    report_type: str = "owner-area-analysis"
    mode: str = "urban"
    scenario: str = "base"
    start_year: int = 2026
    end_year: int = 2035
    avg_kwh_per_session: float = 35.0
    price_per_kwh: float = 7.9
    electricity_cost_per_kwh: float = 4.0
    cpo_gp_rate: float = 0.08
    annual_o_and_m: float = 36_000.0
    project_capex_ex_vat: float = 1_000_000.0
    perception_factor: float = 0.85
    recommended_spec: str | None = None
    metric_label: str | None = None
    metric_value: float | None = None
    owner_gp_per_kwh: float = 0.25
    note: str | None = None


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[^\w\u0E00-\u0E7Fa-zA-Z0-9\- ]+", "", value).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned[:80] or "owner_area_analysis"


def _report_suffix(report_type: str) -> str:
    suffix_map = {
        "owner-area-analysis": "owner_area_analysis",
        "owner-gp-opportunity": "owner_gp_opportunity",
        "investor-case": "investor_case",
    }
    return suffix_map.get(report_type, "owner_area_analysis")


def _fmt_num(value: float | int | None, digits: int = 1) -> str:
    if value is None:
        return "-"
    return f"{float(value):,.{digits}f}"


def _fmt_int(value: float | int | None) -> str:
    if value is None:
        return "-"
    return f"{round(float(value)):,}"


def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_border(cell, color: str = "D2D8E0", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_table_widths(table, widths_inches: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def _style_doc(doc: Document, footer_label: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 22, BLACK),
        ("Heading 1", 15, ACCENT),
        ("Heading 2", 12.5, ACCENT),
        ("Heading 3", 11.5, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TH-EVI | Owner Area Analysis")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(footer_label)
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def _top_name(primary: list[dict[str, Any]], fallback: str) -> str:
    if primary:
        name = str(primary[0].get("name") or "").strip()
        if name:
            return name
    return fallback


def _choose_site_name(req: OwnerAreaReportRequest, first_year: dict[str, Any]) -> str:
    explicit = (req.site_name or "").strip()
    if explicit:
        return explicit
    for key in ("top_pois", "top_business_areas", "top_zones", "top_districts"):
        source = first_year.get(key) or []
        if source:
            name = str(source[0].get("name") or "").strip()
            if name:
                return name
    return f"{req.province} {req.lat:.4f}, {req.lon:.4f}"


def _projection_rows(req: OwnerAreaReportRequest) -> list[dict[str, Any]]:
    from .spatial import analyze_click_location

    rows: list[dict[str, Any]] = []
    for year in range(req.start_year, req.end_year + 1):
        result = analyze_click_location(
            lat=req.lat,
            lon=req.lon,
            province=req.province,
            year=year,
            scenario=req.scenario,
            mode=req.mode,
            avg_kwh_per_session=req.avg_kwh_per_session,
            price_per_kwh=req.price_per_kwh,
        )
        annual_kwh = result["daily_kwh"] * 365.0
        annual_revenue = annual_kwh * req.price_per_kwh
        annual_owner_gp = annual_kwh * req.owner_gp_per_kwh
        annual_cpo_gp = annual_revenue * req.cpo_gp_rate
        annual_electricity_cost = annual_kwh * req.electricity_cost_per_kwh
        annual_investor_cf = annual_revenue - annual_cpo_gp - annual_electricity_cost - annual_owner_gp - req.annual_o_and_m
        result["annual_kwh"] = round(annual_kwh, 1)
        result["annual_revenue"] = round(annual_revenue, 0)
        result["annual_owner_gp"] = round(annual_owner_gp, 0)
        result["annual_cpo_gp"] = round(annual_cpo_gp, 0)
        result["annual_electricity_cost"] = round(annual_electricity_cost, 0)
        result["annual_investor_cf"] = round(annual_investor_cf, 0)
        rows.append(result)
    return rows


def _float_or_none(value: Any) -> float | None:
    try:
        if value in (None, ""):
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _png_chunk(chunk_type: bytes, data: bytes) -> bytes:
    return (
        struct.pack(">I", len(data))
        + chunk_type
        + data
        + struct.pack(">I", zlib.crc32(chunk_type + data) & 0xFFFFFFFF)
    )


def _write_png(path: Path, pixels: list[list[tuple[int, int, int]]]) -> None:
    height = len(pixels)
    width = len(pixels[0]) if height else 0
    raw = bytearray()
    for row in pixels:
        raw.append(0)
        for r, g, b in row:
            raw.extend((r, g, b))
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    content = b"\x89PNG\r\n\x1a\n" + _png_chunk(b"IHDR", ihdr) + _png_chunk(b"IDAT", zlib.compress(bytes(raw), 9)) + _png_chunk(b"IEND", b"")
    path.write_bytes(content)


def _put_pixel(pixels: list[list[tuple[int, int, int]]], x: int, y: int, color: tuple[int, int, int]) -> None:
    if 0 <= y < len(pixels) and 0 <= x < len(pixels[0]):
        pixels[y][x] = color


def _draw_rect(pixels, x0, y0, x1, y1, color, fill=True):
    x0, x1 = sorted((int(x0), int(x1)))
    y0, y1 = sorted((int(y0), int(y1)))
    for y in range(y0, y1 + 1):
        for x in range(x0, x1 + 1):
            if fill or y in (y0, y1) or x in (x0, x1):
                _put_pixel(pixels, x, y, color)


def _draw_line(pixels, x0, y0, x1, y1, color):
    x0, y0, x1, y1 = int(x0), int(y0), int(x1), int(y1)
    dx = abs(x1 - x0)
    sx = 1 if x0 < x1 else -1
    dy = -abs(y1 - y0)
    sy = 1 if y0 < y1 else -1
    err = dx + dy
    while True:
        _put_pixel(pixels, x0, y0, color)
        if x0 == x1 and y0 == y1:
            break
        e2 = 2 * err
        if e2 >= dy:
            err += dy
            x0 += sx
        if e2 <= dx:
            err += dx
            y0 += sy


def _draw_circle(pixels, cx, cy, radius, color, fill=True):
    cx, cy, radius = int(cx), int(cy), int(radius)
    r2 = radius * radius
    for y in range(cy - radius, cy + radius + 1):
        for x in range(cx - radius, cx + radius + 1):
            dx = x - cx
            dy = y - cy
            d2 = dx * dx + dy * dy
            if fill:
                if d2 <= r2:
                    _put_pixel(pixels, x, y, color)
            else:
                if r2 - radius <= d2 <= r2 + radius:
                    _put_pixel(pixels, x, y, color)


def _mix_color(c1: tuple[int, int, int], c2: tuple[int, int, int], ratio: float) -> tuple[int, int, int]:
    return tuple(int(round(a + (b - a) * ratio)) for a, b in zip(c1, c2))


def _heat_color(intensity: float) -> tuple[int, int, int]:
    intensity = max(0.30, min(1.0, float(intensity)))
    for idx in range(1, len(HEAT_COLORS)):
        prev_at, prev_color = HEAT_COLORS[idx - 1]
        next_at, next_color = HEAT_COLORS[idx]
        if intensity <= next_at:
            ratio = (intensity - prev_at) / max(next_at - prev_at, 0.0001)
            return _mix_color(prev_color, next_color, ratio)
    return HEAT_COLORS[-1][1]


def _pil_image_modules():
    try:
        from PIL import Image, ImageDraw
    except Exception as exc:  # pragma: no cover - fallback path is tested elsewhere
        raise RuntimeError("Pillow is not available") from exc
    return Image, ImageDraw


def _latlon_to_tile_xy(lat: float, lon: float, zoom: int) -> tuple[float, float]:
    lat = max(min(lat, 85.05112878), -85.05112878)
    lat_rad = math.radians(lat)
    scale = 2**zoom
    x = (lon + 180.0) / 360.0 * scale
    y = (1.0 - math.asinh(math.tan(lat_rad)) / math.pi) / 2.0 * scale
    return x, y


def _tile_basemap_image(
    lat_min: float,
    lat_max: float,
    lon_min: float,
    lon_max: float,
    zoom: int,
    width: int,
    height: int,
):
    Image, _ = _pil_image_modules()

    nw_x, nw_y = _latlon_to_tile_xy(lat_max, lon_min, zoom)
    se_x, se_y = _latlon_to_tile_xy(lat_min, lon_max, zoom)
    min_tile_x = math.floor(nw_x)
    max_tile_x = math.floor(se_x)
    min_tile_y = math.floor(nw_y)
    max_tile_y = math.floor(se_y)
    tile_count = (max_tile_x - min_tile_x + 1) * (max_tile_y - min_tile_y + 1)
    if tile_count <= 0 or tile_count > 24:
        raise RuntimeError("Too many tiles required for basemap render")

    canvas = Image.new("RGB", ((max_tile_x - min_tile_x + 1) * 256, (max_tile_y - min_tile_y + 1) * 256), MAP_BG)
    headers = {"User-Agent": OSM_USER_AGENT}
    for tile_x in range(min_tile_x, max_tile_x + 1):
        for tile_y in range(min_tile_y, max_tile_y + 1):
            url = OSM_TILE_TEMPLATE.format(z=zoom, x=tile_x, y=tile_y)
            with urllib.request.urlopen(urllib.request.Request(url, headers=headers), timeout=12) as response:
                tile_bytes = response.read()
            tile_image = Image.open(io.BytesIO(tile_bytes)).convert("RGB")
            canvas.paste(tile_image, ((tile_x - min_tile_x) * 256, (tile_y - min_tile_y) * 256))

    left = int(round((nw_x - min_tile_x) * 256))
    top = int(round((nw_y - min_tile_y) * 256))
    right = int(round((se_x - min_tile_x) * 256))
    bottom = int(round((se_y - min_tile_y) * 256))
    right = max(right, left + 32)
    bottom = max(bottom, top + 32)
    cropped = canvas.crop((left, top, right, bottom))
    resized = cropped.resize((width, height))
    return resized


def _tile_count_for_bounds(lat_min: float, lat_max: float, lon_min: float, lon_max: float, zoom: int) -> int:
    nw_x, nw_y = _latlon_to_tile_xy(lat_max, lon_min, zoom)
    se_x, se_y = _latlon_to_tile_xy(lat_min, lon_max, zoom)
    min_tile_x = math.floor(nw_x)
    max_tile_x = math.floor(se_x)
    min_tile_y = math.floor(nw_y)
    max_tile_y = math.floor(se_y)
    return (max_tile_x - min_tile_x + 1) * (max_tile_y - min_tile_y + 1)


def _choose_tile_zoom(
    lat_min: float,
    lat_max: float,
    lon_min: float,
    lon_max: float,
    preferred_zoom: int,
    min_zoom: int = 11,
    max_tiles: int = 24,
) -> int:
    for zoom in range(preferred_zoom, min_zoom - 1, -1):
        if _tile_count_for_bounds(lat_min, lat_max, lon_min, lon_max, zoom) <= max_tiles:
            return zoom
    return min_zoom


def _render_context_map_with_tiles(req: OwnerAreaReportRequest, first_year: dict[str, Any], file_stem: str) -> Path:
    Image, ImageDraw = _pil_image_modules()
    width, height = 900, 560
    margin = 42
    features = _build_map_features(req, first_year)
    lats = [float(item["lat"]) for item in features]
    lons = [float(item["lon"]) for item in features]
    lat_min, lat_max = min(lats), max(lats)
    lon_min, lon_max = min(lons), max(lons)
    lat_pad = max((lat_max - lat_min) * 0.28, 0.012)
    lon_pad = max((lon_max - lon_min) * 0.28, 0.012)
    lat_min -= lat_pad
    lat_max += lat_pad
    lon_min -= lon_pad
    lon_max += lon_pad

    zoom = _choose_tile_zoom(lat_min, lat_max, lon_min, lon_max, preferred_zoom=13, min_zoom=11, max_tiles=24)
    image = _tile_basemap_image(lat_min, lat_max, lon_min, lon_max, zoom=zoom, width=width, height=height).convert("RGBA")
    overlay = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)

    def project(lat: float, lon: float) -> tuple[int, int]:
        x = margin + (lon - lon_min) / max(lon_max - lon_min, 1e-6) * (width - margin * 2)
        y = height - margin - (lat - lat_min) / max(lat_max - lat_min, 1e-6) * (height - margin * 2)
        return int(x), int(y)

    sx, sy = project(req.lat, req.lon)
    color_map = {
        "site": MAP_SITE,
        "poi": MAP_POI,
        "competitor": MAP_COMPETITOR,
        "zone": MAP_ZONE,
        "business": MAP_BUSINESS,
    }
    radius_map = {"site": 10, "poi": 7, "competitor": 7, "zone": 7, "business": 7}

    for feature in features:
        x, y = project(float(feature["lat"]), float(feature["lon"]))
        color = color_map.get(str(feature["kind"]), MAP_POI)
        radius = radius_map.get(str(feature["kind"]), 7)
        if feature["kind"] != "site":
            draw.line((sx, sy, x, y), fill=(255, 255, 255, 180), width=3)
            draw.line((sx, sy, x, y), fill=(120, 136, 153, 150), width=1)
        draw.ellipse((x - radius - 2, y - radius - 2, x + radius + 2, y + radius + 2), fill=(255, 255, 255, 235))
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=(*color, 240), outline=(255, 255, 255, 255))

    label_box = (14, 14, 232, 86)
    draw.rounded_rectangle(label_box, radius=14, fill=(255, 255, 255, 225), outline=(210, 220, 230, 255), width=1)
    image = Image.alpha_composite(image, overlay)
    out_path = REPORT_OUTPUT_DIR / f"{file_stem}_map.png"
    image.convert("RGB").save(out_path, format="PNG")
    return out_path


def _render_heatmap_with_tiles(req: OwnerAreaReportRequest, file_stem: str) -> Path:
    from .heatmap import generate_province_heatmap

    Image, ImageDraw = _pil_image_modules()
    width, height = 900, 560
    margin = 42
    heatmap = generate_province_heatmap(
        req.province,
        year=req.start_year,
        scenario=req.scenario,
        resolution_km=0.5,
        mode=req.mode,
    )
    points = []
    for point in heatmap.get("points", []):
        if abs(float(point["lat"]) - req.lat) <= 0.08 and abs(float(point["lon"]) - req.lon) <= 0.08:
            points.append(point)
    if not points:
        points = heatmap.get("points", [])[:]

    lats = [float(p["lat"]) for p in points] + [req.lat]
    lons = [float(p["lon"]) for p in points] + [req.lon]
    lat_min, lat_max = min(lats), max(lats)
    lon_min, lon_max = min(lons), max(lons)
    lat_pad = max((lat_max - lat_min) * 0.24, 0.01)
    lon_pad = max((lon_max - lon_min) * 0.24, 0.01)
    lat_min -= lat_pad
    lat_max += lat_pad
    lon_min -= lon_pad
    lon_max += lon_pad

    zoom = _choose_tile_zoom(lat_min, lat_max, lon_min, lon_max, preferred_zoom=14, min_zoom=11, max_tiles=24)
    image = _tile_basemap_image(lat_min, lat_max, lon_min, lon_max, zoom=zoom, width=width, height=height).convert("RGBA")
    overlay = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)

    def project(lat: float, lon: float) -> tuple[int, int]:
        x = margin + (lon - lon_min) / max(lon_max - lon_min, 1e-6) * (width - margin * 2)
        y = height - margin - (lat - lat_min) / max(lat_max - lat_min, 1e-6) * (height - margin * 2)
        return int(x), int(y)

    lat_step_deg = float(heatmap.get("lat_step_deg") or 0.0045)
    lon_step_deg = float(heatmap.get("lon_step_deg") or 0.0045)
    for point in points:
        x0, y0 = project(float(point["lat"]) - lat_step_deg / 2, float(point["lon"]) - lon_step_deg / 2)
        x1, y1 = project(float(point["lat"]) + lat_step_deg / 2, float(point["lon"]) + lon_step_deg / 2)
        color = _heat_color(float(point.get("display_intensity") or point.get("intensity") or 0.3))
        fill = (*color, 108)
        outline = (*color, 155)
        draw.rectangle((min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1)), fill=fill, outline=outline, width=1)

    sx, sy = project(req.lat, req.lon)
    draw.ellipse((sx - 12, sy - 12, sx + 12, sy + 12), fill=(255, 255, 255, 235))
    draw.ellipse((sx - 9, sy - 9, sx + 9, sy + 9), fill=(*MAP_SITE, 240))
    draw.ellipse((sx - 15, sy - 15, sx + 15, sy + 15), outline=(*MAP_SITE, 220), width=3)

    image = Image.alpha_composite(image, overlay)
    out_path = REPORT_OUTPUT_DIR / f"{file_stem}_heatmap.png"
    image.convert("RGB").save(out_path, format="PNG")
    return out_path


def _matching_feature(rows: list[dict[str, Any]], name: str, lat_key: str = "lat", lon_key: str = "lon") -> dict[str, Any] | None:
    target = (name or "").strip().lower()
    if not target:
        return None
    for row in rows:
        row_name = str(row.get("name") or row.get("poi_id") or row.get("station_id") or "").strip().lower()
        if row_name == target:
            lat = _float_or_none(row.get(lat_key))
            lon = _float_or_none(row.get(lon_key))
            if lat is not None and lon is not None:
                return row
    return None


def _build_map_features(req: OwnerAreaReportRequest, first_year: dict[str, Any]) -> list[dict[str, Any]]:
    from .spatial import (
        load_business_areas_for_province,
        load_competitors_for_province,
        load_hot_zones_for_province,
        load_pois_for_province,
    )

    features = [{"kind": "site", "name": req.site_name or "Selected site", "lat": req.lat, "lon": req.lon}]
    pois = load_pois_for_province(req.province)
    competitors = load_competitors_for_province(req.province)
    zones = load_hot_zones_for_province(req.province)
    business_areas = load_business_areas_for_province(req.province)

    for item in first_year.get("top_pois", [])[:4]:
        matched = _matching_feature(pois, str(item.get("name") or ""))
        if matched:
            features.append({"kind": "poi", "name": matched.get("name"), "lat": float(matched["lat"]), "lon": float(matched["lon"])})
    for item in first_year.get("top_competitors", [])[:4]:
        matched = _matching_feature(competitors, str(item.get("name") or ""))
        if matched:
            features.append({"kind": "competitor", "name": matched.get("name"), "lat": float(matched["lat"]), "lon": float(matched["lon"])})
    for item in first_year.get("top_zones", [])[:2]:
        matched = _matching_feature(zones, str(item.get("name") or ""), lat_key="center_lat", lon_key="center_lon")
        if matched:
            features.append({"kind": "zone", "name": matched.get("name"), "lat": float(matched["center_lat"]), "lon": float(matched["center_lon"])})
    for item in first_year.get("top_business_areas", [])[:2]:
        matched = _matching_feature(business_areas, str(item.get("name") or ""), lat_key="center_lat", lon_key="center_lon")
        if matched:
            features.append({"kind": "business", "name": matched.get("name"), "lat": float(matched["center_lat"]), "lon": float(matched["center_lon"])})
    return features


def _render_schematic_map(req: OwnerAreaReportRequest, first_year: dict[str, Any], file_stem: str) -> Path:
    try:
        return _render_context_map_with_tiles(req, first_year, file_stem)
    except Exception:
        pass

    width = 900
    height = 560
    margin = 48
    pixels = [[MAP_BG for _ in range(width)] for _ in range(height)]
    features = _build_map_features(req, first_year)
    lats = [float(item["lat"]) for item in features]
    lons = [float(item["lon"]) for item in features]
    lat_min, lat_max = min(lats), max(lats)
    lon_min, lon_max = min(lons), max(lons)
    lat_pad = max((lat_max - lat_min) * 0.25, 0.01)
    lon_pad = max((lon_max - lon_min) * 0.25, 0.01)
    lat_min -= lat_pad
    lat_max += lat_pad
    lon_min -= lon_pad
    lon_max += lon_pad

    def project(lat: float, lon: float) -> tuple[int, int]:
        x = margin + (lon - lon_min) / max(lon_max - lon_min, 1e-6) * (width - margin * 2)
        y = height - margin - (lat - lat_min) / max(lat_max - lat_min, 1e-6) * (height - margin * 2)
        return int(x), int(y)

    _draw_rect(pixels, margin, margin, width - margin, height - margin, MAP_BORDER, fill=False)
    for step in range(1, 5):
        gx = margin + step * (width - margin * 2) / 5
        gy = margin + step * (height - margin * 2) / 5
        _draw_line(pixels, gx, margin, gx, height - margin, MAP_GRID)
        _draw_line(pixels, margin, gy, width - margin, gy, MAP_GRID)

    color_map = {
        "site": MAP_SITE,
        "poi": MAP_POI,
        "competitor": MAP_COMPETITOR,
        "zone": MAP_ZONE,
        "business": MAP_BUSINESS,
    }
    radius_map = {"site": 9, "poi": 7, "competitor": 6, "zone": 6, "business": 6}
    for feature in features:
        x, y = project(float(feature["lat"]), float(feature["lon"]))
        color = color_map.get(str(feature["kind"]), MAP_POI)
        radius = radius_map.get(str(feature["kind"]), 6)
        if feature["kind"] != "site":
            sx, sy = project(req.lat, req.lon)
            _draw_line(pixels, sx, sy, x, y, MAP_GRID)
        _draw_circle(pixels, x, y, radius + 2, (255, 255, 255), fill=True)
        _draw_circle(pixels, x, y, radius, color, fill=True)

    legend_y = height - 26
    legend_x = margin
    legend_items = [("site", MAP_SITE), ("poi", MAP_POI), ("competitor", MAP_COMPETITOR), ("zone", MAP_ZONE), ("business", MAP_BUSINESS)]
    for idx, (_, color) in enumerate(legend_items):
        x0 = legend_x + idx * 34
        _draw_rect(pixels, x0, legend_y, x0 + 18, legend_y + 12, color, fill=True)
        _draw_rect(pixels, x0, legend_y, x0 + 18, legend_y + 12, MAP_BORDER, fill=False)

    out_path = REPORT_OUTPUT_DIR / f"{file_stem}_map.png"
    _write_png(out_path, pixels)
    return out_path


def _render_heatmap_snapshot(req: OwnerAreaReportRequest, file_stem: str) -> Path:
    try:
        return _render_heatmap_with_tiles(req, file_stem)
    except Exception:
        pass

    from .heatmap import generate_province_heatmap

    width = 900
    height = 560
    margin = 48
    pixels = [[MAP_BG for _ in range(width)] for _ in range(height)]
    heatmap = generate_province_heatmap(
        req.province,
        year=req.start_year,
        scenario=req.scenario,
        resolution_km=0.5,
        mode=req.mode,
    )
    points = []
    for point in heatmap.get("points", []):
        if abs(float(point["lat"]) - req.lat) <= 0.08 and abs(float(point["lon"]) - req.lon) <= 0.08:
            points.append(point)
    if not points:
        points = heatmap.get("points", [])[:]

    lats = [float(p["lat"]) for p in points] + [req.lat]
    lons = [float(p["lon"]) for p in points] + [req.lon]
    lat_min, lat_max = min(lats), max(lats)
    lon_min, lon_max = min(lons), max(lons)
    lat_pad = max((lat_max - lat_min) * 0.25, 0.01)
    lon_pad = max((lon_max - lon_min) * 0.25, 0.01)
    lat_min -= lat_pad
    lat_max += lat_pad
    lon_min -= lon_pad
    lon_max += lon_pad

    def project(lat: float, lon: float) -> tuple[int, int]:
        x = margin + (lon - lon_min) / max(lon_max - lon_min, 1e-6) * (width - margin * 2)
        y = height - margin - (lat - lat_min) / max(lat_max - lat_min, 1e-6) * (height - margin * 2)
        return int(x), int(y)

    _draw_rect(pixels, margin, margin, width - margin, height - margin, MAP_BORDER, fill=False)
    for step in range(1, 5):
        gx = margin + step * (width - margin * 2) / 5
        gy = margin + step * (height - margin * 2) / 5
        _draw_line(pixels, gx, margin, gx, height - margin, MAP_GRID)
        _draw_line(pixels, margin, gy, width - margin, gy, MAP_GRID)

    lat_step_deg = float(heatmap.get("lat_step_deg") or 0.0045)
    lon_step_deg = float(heatmap.get("lon_step_deg") or 0.0045)
    for point in points:
        x0, y0 = project(float(point["lat"]) - lat_step_deg / 2, float(point["lon"]) - lon_step_deg / 2)
        x1, y1 = project(float(point["lat"]) + lat_step_deg / 2, float(point["lon"]) + lon_step_deg / 2)
        color = _heat_color(float(point.get("display_intensity") or point.get("intensity") or 0.3))
        _draw_rect(pixels, min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1), color, fill=True)
        _draw_rect(pixels, min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1), (255, 255, 255), fill=False)

    sx, sy = project(req.lat, req.lon)
    _draw_circle(pixels, sx, sy, 11, (255, 255, 255), fill=True)
    _draw_circle(pixels, sx, sy, 8, MAP_SITE, fill=True)
    _draw_circle(pixels, sx, sy, 14, MAP_SITE, fill=False)

    legend_y = height - 26
    legend_x = margin
    for idx, intensity in enumerate([0.30, 0.50, 0.70, 0.90]):
        color = _heat_color(intensity)
        x0 = legend_x + idx * 28
        _draw_rect(pixels, x0, legend_y, x0 + 18, legend_y + 12, color, fill=True)
        _draw_rect(pixels, x0, legend_y, x0 + 18, legend_y + 12, MAP_BORDER, fill=False)

    out_path = REPORT_OUTPUT_DIR / f"{file_stem}_heatmap.png"
    _write_png(out_path, pixels)
    return out_path


def _add_title_block(doc: Document, req: OwnerAreaReportRequest, site_name: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis สำหรับเจ้าของสถานที่")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{site_name} | รายงานวิเคราะห์พื้นที่จาก Heat Map และ Click Analysis")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=7, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    _set_table_widths(meta, [1.8, 4.8])
    rows = [
        ("พื้นที่อ้างอิง", site_name),
        ("จังหวัด", req.province),
        ("พิกัด", f"{req.lat:.6f}, {req.lon:.6f}"),
        ("มุมมองการวิเคราะห์", "Owner Area Analysis"),
        ("โหมดแผนที่", req.mode),
        ("สมมติฐานพลังงาน", f"{_fmt_num(req.avg_kwh_per_session)} kWh/คัน"),
        ("สเปกที่ต้องการดู", req.recommended_spec or "ใช้คำแนะนำจากระบบ"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell)
        _set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def _add_executive_summary(
    doc: Document,
    req: OwnerAreaReportRequest,
    site_name: str,
    first_year: dict[str, Any],
    final_year: dict[str, Any],
) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับผู้บริหาร")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    _set_cell_border(cell, color="C9D6E3", size="10")
    _set_cell_shading(cell, "F4F7FA")
    focus_metric = ""
    if req.metric_label and req.metric_value is not None:
        focus_metric = f" จุดที่เลือกมี {req.metric_label} ประมาณ {_fmt_num(req.metric_value)}."
    note_text = f" หมายเหตุจากผู้ใช้: {req.note}" if req.note else ""
    text = (
        f"{site_name} อยู่ในพื้นที่ที่ระบบอ่านเป็น {first_year['location_type']} และให้สถานะ "
        f"{first_year['eligibility_status']} โดยในปี {req.start_year} ระบบประเมินดีมานด์รวมของพื้นที่ไว้ที่ "
        f"{_fmt_num(first_year['gross_area_demand_sessions'])} คัน/วัน ก่อนหักแรงแข่งขันของคู่แข่ง "
        f"{_fmt_num(first_year['competitor_penalty_sessions'])} คัน/วัน เหลือดีมานด์สุทธิที่พื้นที่นี้สามารถ capture ได้ประมาณ "
        f"{_fmt_num(first_year['net_sessions_per_day'])} คัน/วัน หรือราว {_fmt_num(first_year['daily_kwh'])} kWh/วัน."
        f"{focus_metric} จุดนี้มีแรงหนุนหลักจาก POI สำคัญอย่าง {_top_name(first_year['top_pois'], 'POI ในพื้นที่')} "
        f"และ activity field รอบข้าง ขณะที่คู่แข่งที่กดทับหลักคือ {_top_name(first_year['top_competitors'], 'สถานีรอบข้าง')}."
        f" หากอิงแนวโน้มปัจจุบัน ดีมานด์สุทธิในปี {req.end_year} จะขยับไปที่ประมาณ "
        f"{_fmt_num(final_year['net_sessions_per_day'])} คัน/วัน.{note_text}"
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def _add_snapshot_table(doc: Document, req: OwnerAreaReportRequest, site_name: str, first_year: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Snapshot ของพื้นที่")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(table, [2.6, 1.3, 3.2])
    headers = ["รายการ", "ค่า", "คำอธิบาย"]
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = label
        _set_cell_border(cell)
        _set_cell_shading(cell, "EAF2FB")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = INK

    rows = [
        ("ดีมานด์รวมในพื้นที่", _fmt_num(first_year["gross_area_demand_sessions"]), "ก่อนหักผลของคู่แข่ง"),
        ("แรงแข่งขันของคู่แข่ง", _fmt_num(first_year["competitor_penalty_sessions"]), "ดีมานด์ที่คู่แข่งครองไว้แล้ว"),
        ("ดีมานด์สุทธิในพื้นที่", _fmt_num(first_year["net_sessions_per_day"]), "ดีมานด์ที่พื้นที่นี้ยัง capture ได้"),
        ("พลังงานต่อวัน", _fmt_num(first_year["daily_kwh"]), f"คำนวณที่ {_fmt_num(req.avg_kwh_per_session)} kWh/คัน"),
        ("Location type", str(first_year["location_type"]), "ลักษณะพื้นที่ที่โมเดลอ่านออก"),
        ("Eligibility status", str(first_year["eligibility_status"]), str(first_year["eligibility_reason"])),
        ("AADT ที่ใช้", _fmt_int(first_year["aadt_used"]), "ตัวช่วยสะท้อนทราฟฟิกเบื้องต้น"),
        ("ความเชื่อมั่น", str(first_year["confidence"]), "อิงความแน่นของ POI / competitor / context"),
    ]
    for row_idx, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = value
            _set_cell_border(cells[col_idx])
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if row_idx % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "F9FBFD")


def _add_rank_table(
    doc: Document,
    title: str,
    rows: list[dict[str, Any]],
    columns: list[tuple[str, str]],
) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run(title)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [1.0, 2.6, 1.1, 1.1, 1.2]
    _set_table_widths(table, widths[: len(columns)])
    for idx, (_, label) in enumerate(columns):
        cell = table.cell(0, idx)
        cell.text = label
        _set_cell_border(cell)
        _set_cell_shading(cell, "EAF2FB")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = INK

    if not rows:
        cells = table.add_row().cells
        cells[0].text = "-"
        cells[1].text = "ไม่มีข้อมูลเพียงพอ"
        for cell in cells:
            _set_cell_border(cell)
        return

    for index, item in enumerate(rows, start=1):
        cells = table.add_row().cells
        payload = {"rank": str(index), **item}
        for col_idx, (key, _) in enumerate(columns):
            value = payload.get(key)
            if isinstance(value, float):
                if key in {"distance_km", "sessions", "zone_score", "area_score", "max_kw"}:
                    text = _fmt_num(value)
                else:
                    text = str(value)
            else:
                text = str(value if value not in (None, "") else "-")
            cells[col_idx].text = text
            _set_cell_border(cells[col_idx])
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if index % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "F9FBFD")


def _support_area_rows(first_year: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for item in first_year["top_business_areas"][:3]:
        rows.append({
            "name": item.get("name"),
            "confidence": item.get("confidence"),
            "distance_km": item.get("distance_km"),
            "impact_score": item.get("area_score"),
        })
    for item in first_year["top_zones"][:2]:
        rows.append({
            "name": item.get("name"),
            "confidence": item.get("confidence"),
            "distance_km": item.get("distance_km"),
            "impact_score": item.get("zone_score"),
        })
    return rows


def _add_forecast_table(doc: Document, rows: list[dict[str, Any]]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("แนวโน้มดีมานด์สุทธิ 10 ปี")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(table, [0.9, 1.2, 1.4, 1.3, 1.4])
    headers = ["ปี", "Gross", "แข่งขัน", "Net", "kWh/วัน"]
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = label
        _set_cell_border(cell)
        _set_cell_shading(cell, "EAF2FB")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = INK

    for idx, row in enumerate(rows, start=1):
        values = [
            str(row["year"]),
            _fmt_num(row["gross_area_demand_sessions"]),
            _fmt_num(row["competitor_penalty_sessions"]),
            _fmt_num(row["net_sessions_per_day"]),
            _fmt_num(row["daily_kwh"]),
        ]
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cells[col_idx].text = value
            _set_cell_border(cells[col_idx])
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if idx % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "F9FBFD")

    note = doc.add_paragraph()
    run = note.add_run("Gross = ดีมานด์รวมในพื้นที่ก่อนหักคู่แข่ง | Net = ดีมานด์สุทธิหลังหักแรงแข่งขันแล้ว")
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED


def _add_warnings(doc: Document, warnings: list[str]) -> None:
    if not warnings:
        return
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ข้อควรระวังในการตีความ")

    for warning in warnings[:6]:
        para = doc.add_paragraph(style=None)
        para.paragraph_format.left_indent = Inches(0.2)
        run = para.add_run(f"- {warning}")
        run.font.color.rgb = MUTED
        run.font.size = Pt(10)


def _add_generation_note(doc: Document, generated_at: str) -> None:
    doc.add_section(WD_SECTION_START.CONTINUOUS)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(
        "เอกสารฉบับนี้สร้างจากข้อมูล Heat Map และ Click Analysis ของระบบ TH-EVI เพื่อใช้เป็นข้อมูลตั้งต้นในการพิจารณาพื้นที่ "
        "ไม่ใช่เอกสารยืนยันผลตอบแทนการลงทุนขั้นสุดท้าย"
    )
    run.font.size = Pt(9.3)
    run.font.color.rgb = MUTED

    para = doc.add_paragraph()
    run = para.add_run(f"จัดทำเมื่อ {generated_at}")
    run.font.size = Pt(9.3)
    run.font.color.rgb = MUTED


def _add_map_snapshot(doc: Document, map_path: Path) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("ภาพสรุปพื้นที่")

    doc.add_picture(str(map_path), width=Inches(6.45))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("ภาพสรุปพื้นที่อัตโนมัติจากข้อมูล TH-EVI | แดง = จุดวิเคราะห์ | น้ำเงิน = POI | ม่วง = คู่แข่ง | ส้ม = Hot zone | ฟ้า = Business area")
    run.font.size = Pt(9.2)
    run.font.color.rgb = MUTED


def _add_heatmap_snapshot(doc: Document, image_path: Path) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Heat Map Snapshot")

    doc.add_picture(str(image_path), width=Inches(6.45))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("ภาพ Heat Map รอบจุดที่เลือกจากระบบ TH-EVI | วงสีแดง = จุดวิเคราะห์ | สีเข้ม = ความร้อนสูงกว่า")
    run.font.size = Pt(9.2)
    run.font.color.rgb = MUTED


def _add_owner_gp_title_block(doc: Document, req: OwnerAreaReportRequest, site_name: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis + Owner GP Opportunity")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{site_name} | ข้อเสนอรายได้ GP สำหรับเจ้าของพื้นที่")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=7, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    _set_table_widths(meta, [1.8, 4.8])
    rows = [
        ("พื้นที่อ้างอิง", site_name),
        ("จังหวัด", req.province),
        ("พิกัด", f"{req.lat:.6f}, {req.lon:.6f}"),
        ("รูปแบบความร่วมมือ", f"เจ้าของพื้นที่ไม่ลงทุน | รับ GP {_fmt_num(req.owner_gp_per_kwh, 2)} บาท/kWh"),
        ("ข้อเสนอสถานี", req.recommended_spec or "ใช้คำแนะนำจากระบบ"),
        ("ช่วงพยากรณ์", f"{req.start_year}-{req.end_year}"),
        ("ราคาขายอ้างอิง", f"{_fmt_num(req.price_per_kwh, 1)} บาท/kWh"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell)
        _set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def _add_owner_gp_summary(doc: Document, req: OwnerAreaReportRequest, site_name: str, first_year: dict[str, Any], final_year: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับเจ้าของพื้นที่")

    # recalculated later in dedicated table but summarize here with first/final only
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    _set_cell_border(cell, color="C9D6E3", size="10")
    _set_cell_shading(cell, "F4F7FA")
    text = (
        f"{site_name} เป็นจุดที่ระบบมองว่า {first_year['eligibility_status']} และมีดีมานด์รวมในพื้นที่ประมาณ "
        f"{_fmt_num(first_year['gross_area_demand_sessions'])} คัน/วัน โดยหลังหักแรงแข่งขันแล้ว ยังเหลือดีมานด์สุทธิประมาณ "
        f"{_fmt_num(first_year['net_sessions_per_day'])} คัน/วันในปี {req.start_year}. "
        f"ถ้าพื้นที่นี้เลือกทำดีลแบบเจ้าของไม่ลงทุน แต่รับ GP {_fmt_num(req.owner_gp_per_kwh, 2)} บาทต่อหน่วยจากพลังงานที่จำหน่ายได้ "
        f"จะมีโอกาสรับ GP ปีแรกประมาณ {_fmt_int(first_year['annual_owner_gp'])} บาท/ปี "
        f"และในปี {req.end_year} ประมาณ {_fmt_int(final_year['annual_owner_gp'])} บาท/ปี "
        f"โดยสเปกที่เหมาะกับการคุยเชิงพาณิชย์คือ {req.recommended_spec or 'สเปกตามคำแนะนำของระบบ'} "
        f"ซึ่งช่วยให้เจ้าของเห็นภาพรายได้ประจำจาก EV charging โดยไม่ต้องลง CAPEX เอง."
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def _add_owner_gp_forecast(doc: Document, req: OwnerAreaReportRequest, rows: list[dict[str, Any]]) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("แนวโน้มจำนวนรถ รายได้สถานี และ GP ของเจ้าของพื้นที่")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(table, [0.9, 1.1, 1.5, 1.5, 1.5])
    headers = ["ปี", "คัน/วัน", "รายได้/ปี", "GP/ปี", "GP สะสม"]
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = label
        _set_cell_border(cell)
        _set_cell_shading(cell, "EAF2FB")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = INK

    cumulative_gp = 0.0
    for idx, row in enumerate(rows, start=1):
        cumulative_gp += float(row["annual_owner_gp"])
        values = [
            str(row["year"]),
            _fmt_num(row["net_sessions_per_day"]),
            _fmt_int(row["annual_revenue"]),
            _fmt_int(row["annual_owner_gp"]),
            _fmt_int(cumulative_gp),
        ]
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cells[col_idx].text = value
            _set_cell_border(cells[col_idx])
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if idx % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "F9FBFD")

    note = doc.add_paragraph()
    run = note.add_run("รายได้/ปี เป็นรายได้รวมของสถานีจากสมมติฐานราคาขายต่อหน่วย ส่วน GP/ปี เป็นส่วนแบ่งที่เจ้าของพื้นที่จะได้รับตาม kWh ที่ขายได้")
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED


def _investor_projection_rows(req: OwnerAreaReportRequest, rows: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], tuple[int, float] | None]:
    investor_rows: list[dict[str, Any]] = []
    cumulative_cf = 0.0
    payback: tuple[int, float] | None = None
    for row in rows:
        adjusted_cars = float(row["net_sessions_per_day"]) * req.perception_factor
        annual_kwh = adjusted_cars * req.avg_kwh_per_session * 365.0
        annual_revenue = annual_kwh * req.price_per_kwh
        annual_cpo_gp = annual_revenue * req.cpo_gp_rate
        annual_electricity_cost = annual_kwh * req.electricity_cost_per_kwh
        annual_owner_gp = annual_kwh * req.owner_gp_per_kwh
        annual_cf = annual_revenue - annual_cpo_gp - annual_electricity_cost - annual_owner_gp - req.annual_o_and_m
        previous = cumulative_cf
        cumulative_cf += annual_cf
        if payback is None and previous < req.project_capex_ex_vat <= cumulative_cf and annual_cf > 0:
            months = ((req.project_capex_ex_vat - previous) / annual_cf) * 12.0
            payback = (int(row["year"]), months)
        investor_rows.append({
            "year": row["year"],
            "modeled_cars_per_day": row["net_sessions_per_day"],
            "adjusted_cars_per_day": adjusted_cars,
            "annual_kwh": annual_kwh,
            "annual_revenue": annual_revenue,
            "annual_cpo_gp": annual_cpo_gp,
            "annual_electricity_cost": annual_electricity_cost,
            "annual_owner_gp": annual_owner_gp,
            "annual_investor_cf": annual_cf,
            "cumulative_investor_cf": cumulative_cf,
        })
    return investor_rows, payback


def _fmt_payback_timing(payback: tuple[int, float] | None) -> str:
    if not payback:
        return "ยังไม่คืนทุนภายในช่วงพยากรณ์"
    months_th = [
        "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
        "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม",
    ]
    year, months_into_year = payback
    month_index = min(max(int(months_into_year), 0), 11)
    return f"คาดว่าจะคืนทุนประมาณเดือน{months_th[month_index]} {year}"


def _add_investor_title_block(doc: Document, req: OwnerAreaReportRequest, site_name: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Area Analysis + Investor Case")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"{site_name} | มุมมองนักลงทุน")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    meta = doc.add_table(rows=8, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    _set_table_widths(meta, [1.8, 4.8])
    rows = [
        ("พื้นที่อ้างอิง", site_name),
        ("จังหวัด", req.province),
        ("พิกัด", f"{req.lat:.6f}, {req.lon:.6f}"),
        ("รูปแบบเอกสาร", "Investor Case"),
        ("มูลค่าโครงการ", f"{_fmt_int(req.project_capex_ex_vat)} บาท (ไม่รวม VAT)"),
        ("ข้อเสนอสถานี", req.recommended_spec or "ใช้คำแนะนำจากระบบ"),
        ("ช่วงพยากรณ์", f"{req.start_year}-{req.end_year}"),
        ("Perception factor", f"{_fmt_num(req.perception_factor, 2)} ของ demand สุทธิ"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = label
        right.text = value
        for cell in (left, right):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell)
        _set_cell_shading(left, "F2F4F7")
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = INK


def _add_investor_summary(doc: Document, req: OwnerAreaReportRequest, site_name: str, first_year: dict[str, Any], final_year: dict[str, Any], payback: tuple[int, float] | None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("สรุปสำหรับนักลงทุน")

    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = box.cell(0, 0)
    _set_cell_border(cell, color="C9D6E3", size="10")
    _set_cell_shading(cell, "F4F7FA")
    text = (
        f"{site_name} เป็นจุดที่มีดีมานด์สุทธิเริ่มต้นประมาณ {_fmt_num(first_year['modeled_cars_per_day'])} คัน/วัน "
        f"แต่ในมุมนักลงทุน เอกสารฉบับนี้ลดลงด้วย perception factor เหลือประมาณ {_fmt_num(first_year['adjusted_cars_per_day'])} คัน/วัน "
        f"เพื่อสะท้อนการแปลง demand ให้เป็นการใช้งานจริงของสถานี. "
        f"ภายใต้สมมติฐานราคาขาย {_fmt_num(req.price_per_kwh,1)} บาท/kWh, ค่าไฟ {_fmt_num(req.electricity_cost_per_kwh,1)} บาท/kWh, "
        f"CPO GP {_fmt_num(req.cpo_gp_rate * 100,1)}%, GP เจ้าของพื้นที่ {_fmt_num(req.owner_gp_per_kwh,2)} บาท/kWh และ O&M {_fmt_int(req.annual_o_and_m)} บาท/ปี "
        f"โครงการให้กระแสเงินสดปีแรกประมาณ {_fmt_int(first_year['annual_investor_cf'])} บาท/ปี "
        f"และปีสุดท้ายประมาณ {_fmt_int(final_year['annual_investor_cf'])} บาท/ปี "
        f"โดย {_fmt_payback_timing(payback)}"
    )
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.color.rgb = INK
    run.font.size = Pt(10.8)


def _add_investor_forecast(doc: Document, req: OwnerAreaReportRequest, rows: list[dict[str, Any]], payback: tuple[int, float] | None) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run("Forecast กระแสเงินสดและการคืนทุน")

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_widths(table, [0.9, 1.0, 1.0, 1.4, 1.2, 1.2, 1.3, 1.4])
    headers = ["ปี", "โมเดล", "หลังปรับ", "รายได้/ปี", "CPO/ปี", "ค่าไฟ/ปี", "CF/ปี", "CF สะสม"]
    for idx, label in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = label
        _set_cell_border(cell)
        _set_cell_shading(cell, "EAF2FB")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = INK

    for idx, row in enumerate(rows, start=1):
        values = [
            str(row["year"]),
            _fmt_num(row["modeled_cars_per_day"]),
            _fmt_num(row["adjusted_cars_per_day"]),
            _fmt_int(row["annual_revenue"]),
            _fmt_int(row["annual_cpo_gp"]),
            _fmt_int(row["annual_electricity_cost"]),
            _fmt_int(row["annual_investor_cf"]),
            _fmt_int(row["cumulative_investor_cf"]),
        ]
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cells[col_idx].text = value
            _set_cell_border(cells[col_idx])
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if idx % 2 == 0:
            for cell in cells:
                _set_cell_shading(cell, "F9FBFD")

    note = doc.add_paragraph()
    run = note.add_run(
        f"เงินลงทุนตั้งต้น {_fmt_int(req.project_capex_ex_vat)} บาท | GP เจ้าของพื้นที่ {_fmt_num(req.owner_gp_per_kwh,2)} บาท/kWh | {_fmt_payback_timing(payback)}"
    )
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED


def _report_font_candidates() -> list[Path]:
    windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
    fonts_dir = windir / "Fonts"
    return [
        fonts_dir / "THSarabunNew.ttf",
        fonts_dir / "THSarabun.ttf",
        fonts_dir / "LeelawUI.ttf",
        fonts_dir / "tahoma.ttf",
    ]


def _build_pdf_styles():
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "Helvetica"
    bold_font_name = "Helvetica-Bold"
    for font_path in _report_font_candidates():
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("THEVIRegular", str(font_path)))
            pdfmetrics.registerFont(TTFont("THEVIBold", str(font_path)))
            font_name = "THEVIRegular"
            bold_font_name = "THEVIBold"
            break

    sample = getSampleStyleSheet()
    styles = {
        "font": font_name,
        "font_bold": bold_font_name,
        "title": ParagraphStyle(
            "THTitle",
            parent=sample["Title"],
            fontName=bold_font_name,
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#000000"),
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "THSubtitle",
            parent=sample["Normal"],
            fontName=font_name,
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#59636E"),
            spaceAfter=10,
        ),
        "heading": ParagraphStyle(
            "THHeading",
            parent=sample["Heading1"],
            fontName=bold_font_name,
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=8,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "THBody",
            parent=sample["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=13,
            textColor=colors.HexColor("#0B2545"),
        ),
        "small": ParagraphStyle(
            "THSmall",
            parent=sample["BodyText"],
            fontName=font_name,
            fontSize=8.8,
            leading=11,
            textColor=colors.HexColor("#59636E"),
        ),
        "cell": ParagraphStyle(
            "THCell",
            parent=sample["BodyText"],
            fontName=font_name,
            fontSize=9.2,
            leading=11.2,
            textColor=colors.HexColor("#0B2545"),
        ),
        "cell_bold": ParagraphStyle(
            "THCellBold",
            parent=sample["BodyText"],
            fontName=bold_font_name,
            fontSize=9.2,
            leading=11.2,
            textColor=colors.HexColor("#0B2545"),
        ),
    }
    return styles


def _pdf_paragraph(text: str, style, escape: bool = True):
    from reportlab.platypus import Paragraph

    content = html.escape(text).replace("\n", "<br/>") if escape else text
    return Paragraph(content, style)


def _pdf_table(rows: list[list[Any]], col_widths: list[float] | None = None, repeat_rows: int = 1):
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle

    widths = [value * inch for value in col_widths] if col_widths else None
    table = Table(rows, colWidths=widths, repeatRows=repeat_rows, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAF2FB")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#D2D8E0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEADING", (0, 0), (-1, -1), 11),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def _pdf_box(text: str, styles):
    rows = [[_pdf_paragraph(text, styles["body"])]]
    table = _pdf_table(rows, [6.55], repeat_rows=0)
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F7FA")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#C9D6E3")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    return table


def _build_pdf_story(
    req: OwnerAreaReportRequest,
    projection_rows: list[dict[str, Any]],
    first_year: dict[str, Any],
    final_year: dict[str, Any],
    site_name: str,
    map_path: Path,
    heatmap_path: Path,
):
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage, Spacer

    styles = _build_pdf_styles()
    story: list[Any] = []
    title_lookup = {
        "owner-area-analysis": "Area Analysis สำหรับเจ้าของสถานที่",
        "owner-gp-opportunity": "Area Analysis + Owner GP Opportunity",
        "investor-case": "Area Analysis + Investor Case",
    }
    subtitle_lookup = {
        "owner-area-analysis": f"{site_name} | รายงานวิเคราะห์พื้นที่จาก Heat Map และ Click Analysis",
        "owner-gp-opportunity": f"{site_name} | ข้อเสนอรายได้ GP สำหรับเจ้าของพื้นที่",
        "investor-case": f"{site_name} | มุมมองนักลงทุน",
    }
    story.append(_pdf_paragraph(title_lookup.get(req.report_type, title_lookup["owner-area-analysis"]), styles["title"]))
    story.append(_pdf_paragraph(subtitle_lookup.get(req.report_type, subtitle_lookup["owner-area-analysis"]), styles["subtitle"]))

    meta_rows = [
        [_pdf_paragraph("พื้นที่อ้างอิง", styles["cell_bold"]), _pdf_paragraph(site_name, styles["cell"])],
        [_pdf_paragraph("จังหวัด", styles["cell_bold"]), _pdf_paragraph(req.province, styles["cell"])],
        [_pdf_paragraph("พิกัด", styles["cell_bold"]), _pdf_paragraph(f"{req.lat:.6f}, {req.lon:.6f}", styles["cell"])],
        [_pdf_paragraph("โหมดแผนที่", styles["cell_bold"]), _pdf_paragraph(req.mode, styles["cell"])],
        [_pdf_paragraph("ช่วงพยากรณ์", styles["cell_bold"]), _pdf_paragraph(f"{req.start_year}-{req.end_year}", styles["cell"])],
        [_pdf_paragraph("ข้อเสนอสถานี", styles["cell_bold"]), _pdf_paragraph(req.recommended_spec or "ใช้คำแนะนำจากระบบ", styles["cell"])],
    ]
    if req.report_type == "owner-gp-opportunity":
        meta_rows.append([_pdf_paragraph("GP เจ้าของพื้นที่", styles["cell_bold"]), _pdf_paragraph(f"{_fmt_num(req.owner_gp_per_kwh, 2)} บาท/kWh", styles["cell"])])
    if req.report_type == "investor-case":
        meta_rows.append([_pdf_paragraph("มูลค่าโครงการ", styles["cell_bold"]), _pdf_paragraph(f"{_fmt_int(req.project_capex_ex_vat)} บาท (ไม่รวม VAT)", styles["cell"])])
        meta_rows.append([_pdf_paragraph("Perception factor", styles["cell_bold"]), _pdf_paragraph(_fmt_num(req.perception_factor, 2), styles["cell"])])
    story.append(_pdf_table(meta_rows, [1.8, 4.8], repeat_rows=0))
    story.append(Spacer(1, 0.14 * inch))

    focus_metric = ""
    if req.metric_label and req.metric_value is not None:
        focus_metric = f" จุดที่เลือกมี {req.metric_label} ประมาณ {_fmt_num(req.metric_value)}."
    note_text = f" หมายเหตุจากผู้ใช้: {req.note}" if req.note else ""
    if req.report_type == "owner-gp-opportunity":
        summary_text = (
            f"{site_name} เป็นจุดที่ระบบมองว่า {first_year['eligibility_status']} และมีดีมานด์รวมในพื้นที่ประมาณ "
            f"{_fmt_num(first_year['gross_area_demand_sessions'])} คัน/วัน โดยหลังหักแรงแข่งขันแล้ว ยังเหลือดีมานด์สุทธิประมาณ "
            f"{_fmt_num(first_year['net_sessions_per_day'])} คัน/วันในปี {req.start_year}. "
            f"หากเจ้าของพื้นที่เลือกทำดีลแบบไม่ลงทุนเอง แต่รับ GP {_fmt_num(req.owner_gp_per_kwh, 2)} บาทต่อหน่วยจากพลังงานที่จำหน่ายได้ "
            f"จะมีโอกาสรับ GP ปีแรกประมาณ {_fmt_int(first_year['annual_owner_gp'])} บาท/ปี และในปี {req.end_year} ประมาณ "
            f"{_fmt_int(final_year['annual_owner_gp'])} บาท/ปี."
        )
    elif req.report_type == "investor-case":
        investor_rows, payback = _investor_projection_rows(req, projection_rows)
        summary_text = (
            f"{site_name} มีดีมานด์สุทธิจากโมเดลเริ่มต้นประมาณ {_fmt_num(investor_rows[0]['modeled_cars_per_day'])} คัน/วัน "
            f"และหลังปรับด้วย perception factor เหลือประมาณ {_fmt_num(investor_rows[0]['adjusted_cars_per_day'])} คัน/วัน. "
            f"ภายใต้สมมติฐานราคาขาย {_fmt_num(req.price_per_kwh,1)} บาท/kWh, ค่าไฟ {_fmt_num(req.electricity_cost_per_kwh,1)} บาท/kWh, "
            f"CPO GP {_fmt_num(req.cpo_gp_rate * 100,1)}%, GP เจ้าของพื้นที่ {_fmt_num(req.owner_gp_per_kwh,2)} บาท/kWh และ O&M {_fmt_int(req.annual_o_and_m)} บาท/ปี "
            f"โครงการให้กระแสเงินสดปีแรกประมาณ {_fmt_int(investor_rows[0]['annual_investor_cf'])} บาท/ปี และ {_fmt_payback_timing(payback)}."
        )
    else:
        summary_text = (
            f"{site_name} อยู่ในพื้นที่ที่ระบบอ่านเป็น {first_year['location_type']} และให้สถานะ {first_year['eligibility_status']} "
            f"โดยในปี {req.start_year} ระบบประเมินดีมานด์รวมของพื้นที่ไว้ที่ {_fmt_num(first_year['gross_area_demand_sessions'])} คัน/วัน "
            f"ก่อนหักแรงแข่งขันของคู่แข่ง {_fmt_num(first_year['competitor_penalty_sessions'])} คัน/วัน เหลือดีมานด์สุทธิประมาณ "
            f"{_fmt_num(first_year['net_sessions_per_day'])} คัน/วัน หรือราว {_fmt_num(first_year['daily_kwh'])} kWh/วัน."
            f"{focus_metric}{note_text}"
        )
    story.append(_pdf_paragraph("สรุปสำหรับผู้บริหาร", styles["heading"]))
    story.append(_pdf_box(summary_text, styles))
    story.append(Spacer(1, 0.12 * inch))

    story.append(_pdf_paragraph("Snapshot ของพื้นที่", styles["heading"]))
    snapshot_rows = [
        [_pdf_paragraph("รายการ", styles["cell_bold"]), _pdf_paragraph("ค่า", styles["cell_bold"]), _pdf_paragraph("คำอธิบาย", styles["cell_bold"])],
        [_pdf_paragraph("ดีมานด์รวมในพื้นที่", styles["cell"]), _pdf_paragraph(_fmt_num(first_year["gross_area_demand_sessions"]), styles["cell"]), _pdf_paragraph("ก่อนหักผลของคู่แข่ง", styles["cell"])],
        [_pdf_paragraph("แรงแข่งขันของคู่แข่ง", styles["cell"]), _pdf_paragraph(_fmt_num(first_year["competitor_penalty_sessions"]), styles["cell"]), _pdf_paragraph("ดีมานด์ที่คู่แข่งครองไว้แล้ว", styles["cell"])],
        [_pdf_paragraph("ดีมานด์สุทธิในพื้นที่", styles["cell"]), _pdf_paragraph(_fmt_num(first_year["net_sessions_per_day"]), styles["cell"]), _pdf_paragraph("ดีมานด์ที่พื้นที่นี้ยัง capture ได้", styles["cell"])],
        [_pdf_paragraph("พลังงานต่อวัน", styles["cell"]), _pdf_paragraph(_fmt_num(first_year["daily_kwh"]), styles["cell"]), _pdf_paragraph(f"คำนวณที่ {_fmt_num(req.avg_kwh_per_session)} kWh/คัน", styles["cell"])],
        [_pdf_paragraph("Location type", styles["cell"]), _pdf_paragraph(str(first_year["location_type"]), styles["cell"]), _pdf_paragraph("ลักษณะพื้นที่ที่โมเดลอ่านออก", styles["cell"])],
        [_pdf_paragraph("Eligibility status", styles["cell"]), _pdf_paragraph(str(first_year["eligibility_status"]), styles["cell"]), _pdf_paragraph(str(first_year["eligibility_reason"]), styles["cell"])],
    ]
    story.append(_pdf_table(snapshot_rows, [2.6, 1.3, 3.2]))
    story.append(Spacer(1, 0.12 * inch))

    def rank_rows(title: str, rows: list[dict[str, Any]], columns: list[tuple[str, str]]):
        story.append(_pdf_paragraph(title, styles["heading"]))
        table_rows = [[_pdf_paragraph(label, styles["cell_bold"]) for _, label in columns]]
        if not rows:
            table_rows.append([_pdf_paragraph("-", styles["cell"]), _pdf_paragraph("ไม่มีข้อมูลเพียงพอ", styles["cell"])] + [_pdf_paragraph("", styles["cell"]) for _ in range(len(columns) - 2)])
        else:
            for index, item in enumerate(rows, start=1):
                payload = {"rank": str(index), **item}
                line = []
                for key, _ in columns:
                    value = payload.get(key)
                    if isinstance(value, float):
                        text = _fmt_num(value)
                    else:
                        text = str(value if value not in (None, "") else "-")
                    line.append(_pdf_paragraph(text, styles["cell"]))
                table_rows.append(line)
        story.append(_pdf_table(table_rows, [1.0, 2.6, 1.1, 1.1, 1.2]))
        story.append(Spacer(1, 0.12 * inch))

    rank_rows(
        "POI สำคัญที่หนุนพื้นที่",
        first_year["top_pois"],
        [("rank", "#"), ("name", "ชื่อ POI"), ("category", "ประเภท"), ("distance_km", "ระยะ กม."), ("sessions", "แรงหนุน")],
    )
    rank_rows(
        "คู่แข่งหลักในพื้นที่",
        first_year["top_competitors"],
        [("rank", "#"), ("name", "ชื่อสถานี"), ("network", "เครือข่าย"), ("distance_km", "ระยะ กม."), ("sessions", "แรงกด")],
    )
    if req.report_type == "owner-area-analysis":
        rank_rows(
            "โซน/พื้นที่ธุรกิจที่ช่วยหนุน",
            _support_area_rows(first_year),
            [("rank", "#"), ("name", "ชื่อพื้นที่"), ("confidence", "Confidence"), ("distance_km", "ระยะ กม."), ("impact_score", "Score")],
        )

    story.append(_pdf_paragraph("ภาพสรุปพื้นที่", styles["heading"]))
    story.append(RLImage(str(map_path), width=6.2 * inch, height=3.85 * inch))
    story.append(_pdf_paragraph("ภาพสรุปพื้นที่อัตโนมัติจากข้อมูล TH-EVI", styles["small"]))
    story.append(Spacer(1, 0.10 * inch))
    story.append(_pdf_paragraph("Heat Map Snapshot", styles["heading"]))
    story.append(RLImage(str(heatmap_path), width=6.2 * inch, height=3.85 * inch))
    story.append(_pdf_paragraph("ภาพ Heat Map รอบจุดที่เลือกจากระบบ TH-EVI", styles["small"]))
    story.append(Spacer(1, 0.12 * inch))

    if req.report_type == "owner-gp-opportunity":
        story.append(_pdf_paragraph("แนวโน้มจำนวนรถ รายได้สถานี และ GP ของเจ้าของพื้นที่", styles["heading"]))
        cumulative_gp = 0.0
        rows = [[_pdf_paragraph(label, styles["cell_bold"]) for label in ["ปี", "คัน/วัน", "รายได้/ปี", "GP/ปี", "GP สะสม"]]]
        for row in projection_rows:
            cumulative_gp += float(row["annual_owner_gp"])
            rows.append([
                _pdf_paragraph(str(row["year"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["net_sessions_per_day"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["annual_revenue"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["annual_owner_gp"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(cumulative_gp), styles["cell"]),
            ])
        story.append(_pdf_table(rows, [0.9, 1.1, 1.5, 1.5, 1.5]))
    elif req.report_type == "investor-case":
        investor_rows, payback = _investor_projection_rows(req, projection_rows)
        story.append(_pdf_paragraph("Forecast กระแสเงินสดและการคืนทุน", styles["heading"]))
        rows = [[_pdf_paragraph(label, styles["cell_bold"]) for label in ["ปี", "โมเดล", "หลังปรับ", "รายได้/ปี", "CPO/ปี", "ค่าไฟ/ปี", "CF/ปี", "CF สะสม"]]]
        for row in investor_rows:
            rows.append([
                _pdf_paragraph(str(row["year"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["modeled_cars_per_day"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["adjusted_cars_per_day"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["annual_revenue"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["annual_cpo_gp"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["annual_electricity_cost"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["annual_investor_cf"]), styles["cell"]),
                _pdf_paragraph(_fmt_int(row["cumulative_investor_cf"]), styles["cell"]),
            ])
        story.append(_pdf_table(rows, [0.9, 1.0, 1.0, 1.4, 1.2, 1.2, 1.3, 1.4]))
        story.append(_pdf_paragraph(f"เงินลงทุนตั้งต้น {_fmt_int(req.project_capex_ex_vat)} บาท | {_fmt_payback_timing(payback)}", styles["small"]))
    else:
        story.append(_pdf_paragraph("แนวโน้มดีมานด์สุทธิ 10 ปี", styles["heading"]))
        rows = [[_pdf_paragraph(label, styles["cell_bold"]) for label in ["ปี", "Gross", "แข่งขัน", "Net", "kWh/วัน"]]]
        for row in projection_rows:
            rows.append([
                _pdf_paragraph(str(row["year"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["gross_area_demand_sessions"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["competitor_penalty_sessions"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["net_sessions_per_day"]), styles["cell"]),
                _pdf_paragraph(_fmt_num(row["daily_kwh"]), styles["cell"]),
            ])
        story.append(_pdf_table(rows, [0.9, 1.2, 1.4, 1.3, 1.4]))
        story.append(_pdf_paragraph("Gross = ดีมานด์รวมในพื้นที่ก่อนหักคู่แข่ง | Net = ดีมานด์สุทธิหลังหักแรงแข่งขันแล้ว", styles["small"]))

    warnings = first_year.get("warnings") or []
    if warnings:
        story.append(Spacer(1, 0.12 * inch))
        story.append(_pdf_paragraph("ข้อควรระวังในการตีความ", styles["heading"]))
        for warning in warnings[:6]:
            story.append(_pdf_paragraph(f"- {warning}", styles["small"]))
    story.append(Spacer(1, 0.14 * inch))
    story.append(_pdf_paragraph("เอกสารฉบับนี้สร้างจากข้อมูล Heat Map และ Click Analysis ของระบบ TH-EVI เพื่อใช้เป็นข้อมูลตั้งต้นในการพิจารณาพื้นที่ ไม่ใช่เอกสารยืนยันผลตอบแทนการลงทุนขั้นสุดท้าย", styles["small"]))
    story.append(_pdf_paragraph(f"จัดทำเมื่อ {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["small"]))
    return story


def _prepare_pdf_payload(req: OwnerAreaReportRequest) -> dict[str, Any]:
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    projection_rows = _projection_rows(req)
    first_year = projection_rows[0]
    final_year = projection_rows[-1]
    site_name = _choose_site_name(req, first_year)
    file_stem = _safe_filename(site_name)
    map_path = _render_schematic_map(req, first_year, file_stem)
    heatmap_path = _render_heatmap_snapshot(req, file_stem)
    output_path = REPORT_OUTPUT_DIR / f"{file_stem}_{_report_suffix(req.report_type)}.pdf"
    return {
        "request": asdict(req),
        "projection_rows": projection_rows,
        "first_year": first_year,
        "final_year": final_year,
        "site_name": site_name,
        "map_path": str(map_path),
        "heatmap_path": str(heatmap_path),
        "output_path": str(output_path),
    }


def _render_owner_area_analysis_pdf_from_payload(payload: dict[str, Any]) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate

    req = OwnerAreaReportRequest(**payload["request"])
    projection_rows = payload["projection_rows"]
    first_year = payload["first_year"]
    final_year = payload["final_year"]
    site_name = payload["site_name"]
    map_path = Path(payload["map_path"])
    heatmap_path = Path(payload["heatmap_path"])
    output_path = Path(payload["output_path"])

    story = _build_pdf_story(req, projection_rows, first_year, final_year, site_name, map_path, heatmap_path)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title=site_name,
        author="TH-EVI",
    )
    doc.build(story)
    return output_path


def _create_owner_area_analysis_pdf_internal(req: OwnerAreaReportRequest) -> Path:
    return _render_owner_area_analysis_pdf_from_payload(_prepare_pdf_payload(req))


def _bundled_python_executable() -> Path | None:
    base = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "python"
    candidates = [base / "python.exe", base / "python"]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def create_owner_area_analysis_pdf(req: OwnerAreaReportRequest) -> Path:
    if importlib.util.find_spec("reportlab"):
        return _create_owner_area_analysis_pdf_internal(req)

    bundled_python = _bundled_python_executable()
    if bundled_python is None:
        raise RuntimeError("PDF export is not available in this environment.")

    payload = _prepare_pdf_payload(req)
    env = dict(os.environ)
    env["PYTHONIOENCODING"] = "utf-8"
    env["TH_EVI_REPORT_OUTPUT_DIR"] = str(REPORT_OUTPUT_DIR)
    env["PYTHONPATH"] = str(REPO_ROOT)
    script = (
        "import json, sys; "
        f"sys.path.insert(0, {str(REPO_ROOT)!r}); "
        "from th_evi.owner_area_report import _render_owner_area_analysis_pdf_from_payload; "
        "data=json.loads(open(sys.argv[1], 'r', encoding='utf-8').read()); "
        "path=_render_owner_area_analysis_pdf_from_payload(data); "
        "print(path)"
    )
    payload_fd, payload_name = tempfile.mkstemp(prefix="th_evi_pdf_payload_", suffix=".json")
    os.close(payload_fd)
    payload_file = Path(payload_name)
    payload_file.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    try:
        result = subprocess.run(
            [str(bundled_python), "-c", script, str(payload_file)],
            cwd=str(REPO_ROOT),
            env=env,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=True,
        )
    finally:
        payload_file.unlink(missing_ok=True)
    output = result.stdout.strip().splitlines()[-1].strip()
    output_path = Path(output)
    if not output_path.exists():
        raise RuntimeError("PDF export did not produce an output file.")
    return output_path


def create_owner_area_analysis_report(req: OwnerAreaReportRequest) -> Path:
    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    projection_rows = _projection_rows(req)
    first_year = projection_rows[0]
    final_year = projection_rows[-1]
    site_name = _choose_site_name(req, first_year)
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
    file_stem = _safe_filename(site_name)
    map_path = _render_schematic_map(req, first_year, file_stem)
    heatmap_path = _render_heatmap_snapshot(req, file_stem)

    doc = Document()
    _style_doc(doc, footer_label=site_name)
    if req.report_type == "investor-case":
        investor_rows, payback = _investor_projection_rows(req, projection_rows)
        _add_investor_title_block(doc, req, site_name)
        _add_investor_summary(doc, req, site_name, investor_rows[0], investor_rows[-1], payback)
        _add_snapshot_table(doc, req, site_name, first_year)
        _add_rank_table(
            doc,
            "POI สำคัญที่หนุนพื้นที่",
            first_year["top_pois"],
            [
                ("rank", "#"),
                ("name", "ชื่อ POI"),
                ("category", "ประเภท"),
                ("distance_km", "ระยะ กม."),
                ("sessions", "แรงหนุน"),
            ],
        )
        _add_rank_table(
            doc,
            "คู่แข่งหลักในพื้นที่",
            first_year["top_competitors"],
            [
                ("rank", "#"),
                ("name", "ชื่อสถานี"),
                ("network", "เครือข่าย"),
                ("distance_km", "ระยะ กม."),
                ("sessions", "แรงกด"),
            ],
        )
        _add_map_snapshot(doc, map_path)
        _add_heatmap_snapshot(doc, heatmap_path)
        _add_investor_forecast(doc, req, investor_rows, payback)
    elif req.report_type == "owner-gp-opportunity":
        _add_owner_gp_title_block(doc, req, site_name)
        _add_owner_gp_summary(doc, req, site_name, first_year, final_year)
        _add_snapshot_table(doc, req, site_name, first_year)
        _add_rank_table(
            doc,
            "POI สำคัญที่หนุนพื้นที่",
            first_year["top_pois"],
            [
                ("rank", "#"),
                ("name", "ชื่อ POI"),
                ("category", "ประเภท"),
                ("distance_km", "ระยะ กม."),
                ("sessions", "แรงหนุน"),
            ],
        )
        _add_rank_table(
            doc,
            "คู่แข่งหลักในพื้นที่",
            first_year["top_competitors"],
            [
                ("rank", "#"),
                ("name", "ชื่อสถานี"),
                ("network", "เครือข่าย"),
                ("distance_km", "ระยะ กม."),
                ("sessions", "แรงกด"),
            ],
        )
        _add_map_snapshot(doc, map_path)
        _add_heatmap_snapshot(doc, heatmap_path)
        _add_owner_gp_forecast(doc, req, projection_rows)
    else:
        _add_title_block(doc, req, site_name)
        _add_executive_summary(doc, req, site_name, first_year, final_year)
        _add_snapshot_table(doc, req, site_name, first_year)
        _add_rank_table(
            doc,
            "POI สำคัญที่หนุนพื้นที่",
            first_year["top_pois"],
            [
                ("rank", "#"),
                ("name", "ชื่อ POI"),
                ("category", "ประเภท"),
                ("distance_km", "ระยะ กม."),
                ("sessions", "แรงหนุน"),
            ],
        )
        _add_rank_table(
            doc,
            "คู่แข่งหลักในพื้นที่",
            first_year["top_competitors"],
            [
                ("rank", "#"),
                ("name", "ชื่อสถานี"),
                ("network", "เครือข่าย"),
                ("distance_km", "ระยะ กม."),
                ("sessions", "แรงกด"),
            ],
        )
        _add_rank_table(
            doc,
            "โซน/พื้นที่ธุรกิจที่ช่วยหนุน",
            _support_area_rows(first_year),
            [
                ("rank", "#"),
                ("name", "ชื่อพื้นที่"),
                ("confidence", "Confidence"),
                ("distance_km", "ระยะ กม."),
                ("impact_score", "Score"),
            ],
        )
        _add_map_snapshot(doc, map_path)
        _add_heatmap_snapshot(doc, heatmap_path)
        _add_forecast_table(doc, projection_rows)
    _add_warnings(doc, first_year.get("warnings") or [])
    _add_generation_note(doc, generated_at)

    suffix = _report_suffix(req.report_type)
    output_path = REPORT_OUTPUT_DIR / f"{file_stem}_{suffix}.docx"
    doc.save(output_path)
    return output_path
